#!/usr/bin/env python3
"""
generate_report.py
──────────────────
Reads a Researcher Metrics Excel export and populates a
departmental Word report.

Usage:
    python generate_report.py <input.xlsx> [template.docx] [output.docx]

Defaults:
    template  → University_of_New_South_Wales_Departmental_Report_2026.docx
    output    → Departmental_Report_<YYYY-MM-DD>.docx
"""

import sys, os, re, io, math, warnings
import datetime, zoneinfo

import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from matplotlib import rcParams

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore')

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY        = '#1F3864'
BLUE        = '#1F497D'
HDR_COLOR   = '1F3864'   # table header bg
STRIPE      = 'EAF0F7'
WHITE       = 'FFFFFF'
FONT        = 'Arial'

rcParams['font.family'] = ['Arial', 'Helvetica', 'DejaVu Sans', 'sans-serif']
rcParams['axes.spines.top']   = False
rcParams['axes.spines.right'] = False

# A4 with 1-inch margins → usable width
PAGE_DXA = 9026   # twips

# ─────────────────────────────────────────────────────────────────────────────
# CHART HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def chart_col(labels, values, title, xlabel='', ylabel='', figsize=(7, 3.5), color=BLUE):
    fig, ax = plt.subplots(figsize=figsize)
    x = range(len(labels))
    bars = ax.bar(x, values, color=color, edgecolor='white', width=0.6)
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=8, rotation=30, ha='right')
    ax.set_ylabel(ylabel, fontsize=9)
    ax.set_xlabel(xlabel, fontsize=9)
    ax.set_title(title, fontsize=10, fontweight='bold', color=NAVY, pad=8)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{int(v):,}'))
    for bar in bars:
        h = bar.get_height()
        if h > 0:
            ax.text(bar.get_x() + bar.get_width()/2, h,
                    f'{int(h):,}', ha='center', va='bottom', fontsize=7)
    ax.set_facecolor('#FAFAFA')
    fig.patch.set_facecolor('white')
    fig.tight_layout()
    return _buf(fig)


def chart_hbar(labels, values, title, xlabel='', figsize=(7, 4), color=BLUE):
    fig, ax = plt.subplots(figsize=figsize)
    y = range(len(labels))
    bars = ax.barh(list(y), values, color=color, edgecolor='white', height=0.6)
    ax.set_yticks(list(y))
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel(xlabel, fontsize=9)
    ax.set_title(title, fontsize=10, fontweight='bold', color=NAVY, pad=8)
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{int(v):,}'))
    for bar in bars:
        w = bar.get_width()
        if w > 0:
            ax.text(w, bar.get_y() + bar.get_height()/2,
                    f' {int(w):,}', va='center', fontsize=7)
    ax.set_facecolor('#FAFAFA')
    fig.patch.set_facecolor('white')
    fig.tight_layout()
    return _buf(fig)


# ─────────────────────────────────────────────────────────────────────────────
# STYLE SETUP — force all heading/body styles to Arial
# ─────────────────────────────────────────────────────────────────────────────

def _setup_styles(doc):
    """Override built-in heading and Normal styles to use Arial consistently."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    configs = {
        'Normal':    dict(size=10, bold=False, color=None),
        'Heading 1': dict(size=16, bold=True,  color='1F3864'),
        'Heading 2': dict(size=13, bold=True,  color='1F497D'),
        'Heading 3': dict(size=11, bold=True,  color='434343'),
        'Heading 4': dict(size=10, bold=True,  color='434343'),
    }
    for style_name, cfg in configs.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        # rPr in the style element
        el = style.element
        rPr = el.find('.//' + qn('w:rPr'))
        if rPr is None:
            # find or create pPr first, then add rPr after it
            rPr = OxmlElement('w:rPr')
            el.append(rPr)
        # Clear existing font/size/color
        for tag in ('w:rFonts', 'w:sz', 'w:szCs', 'w:color', 'w:b'):
            for old in rPr.findall(qn(tag)):
                rPr.remove(old)
        # Font
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), FONT)
        rFonts.set(qn('w:hAnsi'), FONT)
        rFonts.set(qn('w:cs'),    FONT)
        rPr.insert(0, rFonts)
        # Size
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(cfg['size'] * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(cfg['size'] * 2)))
        rPr.append(szCs)
        # Bold
        if cfg['bold']:
            b = OxmlElement('w:b')
            rPr.append(b)
        # Color
        if cfg['color']:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), cfg['color'])
            rPr.append(color)
        # Also clear any theme-font attributes from the pPr/pStyle
        for pPr_el in el.findall('.//' + qn('w:pPr')):
            for theme in ('w:outlineLvl',):
                pass  # keep structure, just override fonts


# ─────────────────────────────────────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _add_run(para, text, bold=False, italic=False, size=10,
             color_hex=None, font=FONT):
    run = para.add_run(text)
    run.bold, run.italic = bold, italic
    run.font.name = font
    run.font.size = Pt(size)
    if color_hex:
        r, g, b = int(color_hex[:2],16), int(color_hex[2:4],16), int(color_hex[4:],16)
        run.font.color.rgb = RGBColor(r, g, b)
    return run



def _make_bookmark_id(text):
    """Convert heading text to a clean bookmark ID."""
    return re.sub(r'[^a-zA-Z0-9_]', '_', str(text)).strip('_')[:40]


def _add_bookmark(para, bookmark_id):
    """Attach a Word bookmark to a paragraph for TOC hyperlinks."""
    bm_id_val = str(abs(hash(bookmark_id)) % 100000)
    bm_start  = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'),   bm_id_val)
    bm_start.set(qn('w:name'), bookmark_id)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'),     bm_id_val)
    para._p.insert(0, bm_start)
    para._p.append(bm_end)


def _add_internal_hyperlink(para, text, bookmark_id, bold=False, size=10, color_hex='1F3864'):
    """Add a clickable internal hyperlink to a bookmark in the same document."""
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'),  bookmark_id)
    hl.set(qn('w:history'), '1')
    run_el = OxmlElement('w:r')
    rPr    = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), FONT); rf.set(qn('w:hAnsi'), FONT); rPr.append(rf)
    sz = OxmlElement('w:sz');    sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    col = OxmlElement('w:color'); col.set(qn('w:val'), color_hex);         rPr.append(col)
    u   = OxmlElement('w:u');     u.set(qn('w:val'), 'single');            rPr.append(u)
    if bold: rPr.append(OxmlElement('w:b'))
    run_el.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    run_el.append(t); hl.append(run_el); para._p.append(hl)


def _add_heading(doc, text, level=2):
    """Add a heading with consistent Arial font and a bookmark for TOC links."""
    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    run.font.name = FONT
    # colour and spacing by level
    if level == 1:
        run.font.size = Pt(16); run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after  = Pt(10)
    elif level == 2:
        run.font.size = Pt(13); run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after  = Pt(8)
    else:  # 3+
        run.font.size = Pt(11); run.bold = True
        run.font.color.rgb = RGBColor(0x43, 0x43, 0x43)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    _add_bookmark(p, _make_bookmark_id(text))
    return p


def _add_body(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold, run.italic = bold, italic
    run.font.name = FONT
    run.font.size = Pt(size)
    return p


def _insert_image(doc, buf, width_inches=6.0, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Keep image paragraph with the caption that follows it
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(24)
        # Keep caption together and with the image above
        cp.paragraph_format.keep_with_next = False
        cp.paragraph_format.keep_together  = True
        run = cp.add_run(caption)
        run.italic = True
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def _page_break(doc):
    doc.add_page_break()


# ── Table builder ─────────────────────────────────────────────────────────────

def _keep_row_together(row):
    """Keep this row with the next one — applied to all rows forces the whole table onto one page."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    # cantSplit: don't break the row itself across pages
    cantSplit = OxmlElement('w:cantSplit')
    cantSplit.set(qn('w:val'), 'true')
    trPr.append(cantSplit)
    # keep_with_next at paragraph level for each cell paragraph
    for cell in row.cells:
        for para in cell.paragraphs:
            para.paragraph_format.keep_with_next = True


def _keep_table_on_one_page(tbl):
    """Force the entire table onto a single page by setting keep_with_next on every row."""
    rows = tbl.rows
    for row in rows:
        _keep_row_together(row)
    # Clear keep_with_next on the very last row's cell paragraphs so flow resumes after
    for cell in rows[-1].cells:
        for para in cell.paragraphs:
            para.paragraph_format.keep_with_next = False


def _add_table(doc, df, caption=None, col_widths_pct=None,
               header_color=HDR_COLOR, stripe_color=STRIPE):
    """
    Fully formatted table. col_widths_pct: list summing to 100.
    Uses fixed-layout XML so widths are respected by LibreOffice / Word.
    """
    n_cols = len(df.columns)
    if col_widths_pct is None:
        col_widths_pct = [100 / n_cols] * n_cols

    raw = [int(PAGE_DXA * p / 100) for p in col_widths_pct]
    raw[0] += PAGE_DXA - sum(raw)   # absorb rounding

    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(10)
        cp.paragraph_format.space_after  = Pt(3)
        cp.paragraph_format.keep_with_next = True
        _add_run(cp, caption, bold=True, italic=True, size=9.5,
                 color_hex='1F3864')

    tbl = doc.add_table(rows=1 + len(df), cols=n_cols)
    tbl.style = 'Normal Table'

    # ── Fixed layout + total width ────────────────────────────────────────────
    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tblEl.insert(0, tblPr)
    for ts in tblPr.findall(qn('w:tblStyle')):
        tblPr.remove(ts)
    for tag in ('w:tblW', 'w:tblLayout'):
        for el in tblPr.findall(qn(tag)): tblPr.remove(el)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(PAGE_DXA)); tblW.set(qn('w:type'), 'dxa')
    tblPr.insert(0, tblW)
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)

    # ── Grid ──────────────────────────────────────────────────────────────────
    for old in tblEl.findall(qn('w:tblGrid')): tblEl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in raw:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGrid.append(gc)
    tblEl.insert(list(tblEl).index(tblPr) + 1, tblGrid)

    # ── Cell formatter ────────────────────────────────────────────────────────
    def _fmt(cell, ci, text, bold=False, bg='FFFFFF', border_c='C8D8E8',
             align=WD_ALIGN_PARAGRAPH.LEFT):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Width
        for el in tcPr.findall(qn('w:tcW')): tcPr.remove(el)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(raw[ci])); tcW.set(qn('w:type'), 'dxa')
        tcPr.insert(0, tcW)

        # Borders
        for el in tcPr.findall(qn('w:tcBorders')): tcPr.remove(el)
        borders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0');   b.set(qn('w:color'), border_c)
            borders.append(b)
        tcPr.append(borders)

        # Background
        for el in tcPr.findall(qn('w:shd')): tcPr.remove(el)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg.lstrip('#'))
        tcPr.append(shd)

        # Padding
        for el in tcPr.findall(qn('w:tcMar')): tcPr.remove(el)
        mar = OxmlElement('w:tcMar')
        for side, val in (('top',60),('left',100),('bottom',60),('right',100)):
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), str(val)); m.set(qn('w:type'), 'dxa')
            mar.append(m)
        tcPr.append(mar)

        # Text
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        color_hex = 'FFFFFF' if bold else '1A1A1A'
        _add_run(p, str(text) if text is not None else '',
                 bold=bold, size=9, color_hex=color_hex)

    # ── Header ────────────────────────────────────────────────────────────────
    for ci, col in enumerate(df.columns):
        _fmt(tbl.rows[0].cells[ci], ci, col,
             bold=True, bg=header_color, border_c=header_color)

    # ── Data rows ─────────────────────────────────────────────────────────────
    for ri, (_, row_data) in enumerate(df.iterrows()):
        bg = STRIPE if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_data):
            if pd.isna(val) or val == '' or val is None:
                txt = ''
            elif isinstance(val, float) and val == int(val):
                txt = f'{int(val):,}'
            elif isinstance(val, float):
                txt = f'{val:,.2f}'.rstrip('0').rstrip('.')
            elif isinstance(val, int):
                txt = f'{val:,}'
            else:
                txt = str(val).strip()
            is_numeric_col = ci > 0 and (
                isinstance(val, (int, float)) and not pd.isna(val)
                or txt.lower() == 'n/a'
            )
            align = WD_ALIGN_PARAGRAPH.RIGHT if is_numeric_col else WD_ALIGN_PARAGRAPH.LEFT
            _fmt(tbl.rows[ri+1].cells[ci], ci, txt, bg=bg, align=align)

    _keep_table_on_one_page(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def _add_hyperlink_table(doc, rows, col_headers, col_widths_pct,
                         url_col_idx, caption=None,
                         header_color=HDR_COLOR, stripe_color=STRIPE):
    """
    Like _add_table but one column is rendered as a hyperlink.
    rows: list of dicts with keys matching col_headers.
    url_col_idx: which column index gets hyperlinked (0-based).
    url_col_key: 'url' key in each row dict containing the href.
    """
    n_cols = len(col_headers)
    raw = [int(PAGE_DXA * p / 100) for p in col_widths_pct]
    raw[0] += PAGE_DXA - sum(raw)

    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(10)
        cp.paragraph_format.space_after  = Pt(3)
        cp.paragraph_format.keep_with_next = True
        _add_run(cp, caption, bold=True, italic=True, size=9.5,
                 color_hex='1F3864')

    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Normal Table'

    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tblEl.insert(0, tblPr)
    for ts in tblPr.findall(qn('w:tblStyle')): tblPr.remove(ts)
    for tag in ('w:tblW', 'w:tblLayout'):
        for el in tblPr.findall(qn(tag)): tblPr.remove(el)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(PAGE_DXA)); tblW.set(qn('w:type'), 'dxa')
    tblPr.insert(0, tblW)
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    for old in tblEl.findall(qn('w:tblGrid')): tblEl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in raw:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGrid.append(gc)
    tblEl.insert(list(tblEl).index(tblPr) + 1, tblGrid)

    def _base_fmt(cell, ci, bg='FFFFFF', border_c='C8D8E8'):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for el in tcPr.findall(qn('w:tcW')): tcPr.remove(el)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(raw[ci])); tcW.set(qn('w:type'), 'dxa')
        tcPr.insert(0, tcW)
        for el in tcPr.findall(qn('w:tcBorders')): tcPr.remove(el)
        borders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
            b.set(qn('w:space'),'0'); b.set(qn('w:color'), border_c)
            borders.append(b)
        tcPr.append(borders)
        for el in tcPr.findall(qn('w:shd')): tcPr.remove(el)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), bg.lstrip('#'))
        tcPr.append(shd)
        for el in tcPr.findall(qn('w:tcMar')): tcPr.remove(el)
        mar = OxmlElement('w:tcMar')
        for side, val in (('top',60),('left',100),('bottom',60),('right',100)):
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), str(val)); m.set(qn('w:type'), 'dxa')
            mar.append(m)
        tcPr.append(mar)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        return p

    # Header
    for ci, h in enumerate(col_headers):
        p = _base_fmt(tbl.rows[0].cells[ci], ci,
                      bg=header_color, border_c=header_color)
        _add_run(p, h, bold=True, size=9, color_hex='FFFFFF')

    # Data rows
    for ri, row in enumerate(rows):
        bg = STRIPE if ri % 2 == 0 else WHITE
        for ci, key in enumerate(col_headers):
            val  = row.get(key, '')
            cell = tbl.rows[ri+1].cells[ci]
            p    = _base_fmt(cell, ci, bg=bg)
            if ci == url_col_idx and row.get('_url'):
                _add_hyperlink(p, str(val), row['_url'])
            else:
                txt = '' if (val is None or (isinstance(val, float) and math.isnan(val))) else str(val)
                align = WD_ALIGN_PARAGRAPH.RIGHT \
                    if ci > 0 and isinstance(val, (int, float)) and not pd.isna(val) \
                    else WD_ALIGN_PARAGRAPH.LEFT
                p.alignment = align
                _add_run(p, txt, size=9, color_hex='1A1A1A')

    _keep_table_on_one_page(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def _add_hyperlink(para, display_text, url):
    """Add a hyperlink run to an existing paragraph."""
    part = para.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')
    run_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT); rFonts.set(qn('w:hAnsi'), FONT)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18')   # 9pt
    rPr.append(sz)
    # Colour + underline
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
    rPr.append(u)
    run_el.append(rPr)
    t = OxmlElement('w:t')
    t.text = display_text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    run_el.append(t)
    hyperlink.append(run_el)
    para._p.append(hyperlink)
    return hyperlink


# ─────────────────────────────────────────────────────────────────────────────
# DATA LOADING
# ─────────────────────────────────────────────────────────────────────────────

def load_excel(path):
    xl = pd.ExcelFile(path)
    sheets = xl.sheet_names
    def read(name):
        return pd.read_excel(xl, sheet_name=name) if name in sheets else pd.DataFrame()

    # Auto-detect benchmarking file in same folder as xlsx
    bench_path = None
    base_dir = os.path.dirname(os.path.abspath(path))
    for fname in os.listdir(base_dir):
        if fname.lower().endswith('.xlsx') and 'benchmark' in fname.lower():
            bench_path = os.path.join(base_dir, fname)
            break

    bench = {}
    if bench_path:
        try:
            bxl = pd.ExcelFile(bench_path)

            def _parse_summary(name):
                """Parse a summary sheet whose data starts at row index 2 (0-based)."""
                if name not in bxl.sheet_names:
                    return pd.DataFrame()
                raw = pd.read_excel(bxl, sheet_name=name, header=None)
                # Find header row containing 'Institution'
                hdr = None
                for i, row in raw.iterrows():
                    if any(str(v).strip() == 'Institution' for v in row):
                        hdr = i; break
                if hdr is None:
                    return pd.DataFrame()
                df = raw.iloc[hdr:].copy()
                df.columns = [str(v).strip() for v in df.iloc[0]]
                df = df.iloc[1:].reset_index(drop=True)
                df = df[df['Institution'].notna() &
                        (df['Institution'].astype(str).str.strip() != 'nan')]
                for col in df.columns[1:]:
                    df[col] = pd.to_numeric(
                        df[col].astype(str).str.replace('%','',regex=False),
                        errors='coerce')
                return df.reset_index(drop=True)

            def _parse_top_cited_multiyr(name):
                """
                Parse a multi-year Top Cited sheet (one block per year) and return a
                per-institution aggregate suitable for _build_university_comparison.
                Columns returned: Institution, Publications,
                  Top 1% % of Total, Top 5% % of Total, Top 10% % of Total,
                  Top 1% Global %, Top 5% Global %, Top 10% Global %
                """
                if name not in bxl.sheet_names:
                    return pd.DataFrame()
                raw = pd.read_excel(bxl, sheet_name=name, header=None)

                blocks = []
                i = 0
                while i < len(raw):
                    cell = str(raw.iloc[i, 0]).strip()
                    # Detect a year header row
                    try:
                        yr = int(float(cell))
                        if 2000 <= yr <= 2035:
                            # Scan next few rows for the Institution header
                            for j in range(i+1, min(i+6, len(raw))):
                                if any(str(v).strip() == 'Institution' for v in raw.iloc[j]):
                                    cols = [str(v).strip() for v in raw.iloc[j]]
                                    block_rows = []
                                    for k in range(j+1, len(raw)):
                                        inst = str(raw.iloc[k, 0]).strip()
                                        if inst in ('nan', '', 'NaN'):
                                            break
                                        row_dict = dict(zip(cols, raw.iloc[k].tolist()))
                                        row_dict['_year'] = yr
                                        block_rows.append(row_dict)
                                    if block_rows:
                                        blocks.append(pd.DataFrame(block_rows))
                                    i = j + len(block_rows)
                                    break
                            i += 1
                            continue
                    except (ValueError, TypeError):
                        pass
                    i += 1

                if not blocks:
                    return pd.DataFrame()

                combined = pd.concat(blocks, ignore_index=True)
                # Keep only institution rows (not ALL AUSTRALIA / ALL GLOBAL)
                combined = combined[~combined['Institution'].astype(str).str.startswith('ALL')]

                # Convert pct strings to float.
                # The benchmarking script names national columns dynamically using the
                # detected country code (e.g. 'AU Top 1%', 'NZ Top 1%', 'US Top 1%').
                # Build the pct_map by scanning actual column names so any country works.
                nat_top1_col  = next((c for c in combined.columns
                                      if c.endswith('Top 1%')  and c != 'Global Top 1%'),  None)
                nat_top5_col  = next((c for c in combined.columns
                                      if c.endswith('Top 5%')  and c != 'Global Top 5%'),  None)
                nat_top10_col = next((c for c in combined.columns
                                      if c.endswith('Top 10%') and c != 'Global Top 10%'), None)

                pct_map = {}
                if nat_top1_col:  pct_map[nat_top1_col]  = 'Top 1% % of Total'
                if nat_top5_col:  pct_map[nat_top5_col]  = 'Top 5% % of Total'
                if nat_top10_col: pct_map[nat_top10_col] = 'Top 10% % of Total'
                pct_map['Global Top 1%']  = 'Top 1% Global %'
                pct_map['Global Top 5%']  = 'Top 5% Global %'
                pct_map['Global Top 10%'] = 'Top 10% Global %'

                # If no national columns found, fall back to using Global columns for 'Total' too
                if not nat_top1_col:
                    pct_map['Global Top 1%']  = 'Top 1% % of Total'
                    pct_map['Global Top 5%']  = 'Top 5% % of Total'
                    pct_map['Global Top 10%'] = 'Top 10% % of Total'

                pubs_col = 'Total Pubs (year)'
                combined[pubs_col] = pd.to_numeric(combined.get(pubs_col, pd.Series()), errors='coerce')
                for src, dst in pct_map.items():
                    if src in combined.columns:
                        combined[dst] = pd.to_numeric(
                            combined[src].astype(str).str.replace('%','',regex=False),
                            errors='coerce')

                agg = {'Publications': (pubs_col, 'sum')}
                for dst in pct_map.values():
                    if dst in combined.columns:
                        agg[dst] = (dst, 'mean')

                result = combined.groupby('Institution').agg(**agg).reset_index()
                # Add FCR stub columns expected downstream
                for tier in (1, 5, 10):
                    k = f'Top {tier}% FCR (geo mean)'
                    if k not in result.columns:
                        result[k] = float('nan')
                return result

            # Auto-detect FoR codes from sheet names like 'Summary (XXXX)'
            import re as _re
            detected_fors = sorted(set(
                _re.search(r'Summary \((\d+)\)', s).group(1)
                for s in bxl.sheet_names
                if _re.search(r'Summary \((\d+)\)', s)
            ))
            bench = {
                'summary':   pd.DataFrame(),
                'top_cited': pd.DataFrame(),
                '_path':     bench_path,
                '_fors':     detected_fors,
            }
            for code in detected_fors:
                bench[f'summary_{code}']   = _parse_summary(f'Summary ({code})')
                bench[f'top_cited_{code}'] = _parse_top_cited_multiyr(f'Top Cited ({code})')
            print(f'  FoR codes detected: {detected_fors}')
            print(f'📊  Benchmarking file : {os.path.basename(bench_path)}')
        except Exception as e:
            import traceback; traceback.print_exc()
            print(f'⚠️   Could not load benchmarking file: {e}')

    return {
        'bench': bench,
        'impact':        read('Impact Summary'),
        'pub_summary':   read('Publications Summary'),
        'journals':      read('Journals Summary'),
        'for_':          read('Fields of Research'),
        'collab':        read('Collaboration Summary'),
        'org_collab':    read('Org Type Collaboration'),
        'funders':       read('Funders Summary'),
        'policy':        read('Policy Publishers'),
        'patents':       read('Patents Summary'),
        'clinical':      read('Clinical Trials Summary'),
        'complete':      read('Complete Publications'),
        'report_info':   read('Report Information'),
        'country_collab':    read('Country Collaboration'),
        'top_researchers':   read('Top Researchers by Impact'),
        'corporate_citations': read('Corporate Citations'),
        'industry_coauth':     read('Industry Co-authorship'),
        'datasets':            read('Datasets'),
    }




def get_report_meta(data):
    meta = {}
    try:
        for label, value in data['report_info'].iloc[:, :2].values.tolist():
            if pd.isna(label): continue
            label = str(label).strip()
            if 'Faculty'      in label: meta['faculty']      = str(value).strip()
            if 'Reporting'    in label: meta['period']       = str(value).strip()
            if 'Number'       in label: meta['n_researchers']= str(value).strip()
            if 'Total Pub'    in label: meta['total_pubs']   = str(value).strip()
            if 'Data Source'  in label: meta['data_source']  = str(value).strip()
            if 'Institution'  in label: meta['institution']  = str(value).strip()
    except Exception:
        pass
    # Always use current AEDT datetime (UTC+11)
    aedt = zoneinfo.ZoneInfo('Australia/Sydney')
    now  = datetime.datetime.now(aedt)
    day  = now.day
    suffix = {1:'st',2:'nd',3:'rd'}.get(day if day < 20 else day % 10, 'th')
    meta['date'] = now.strftime(f'{day}{suffix} %B %Y')
    meta['year'] = now.strftime('%Y')
    return meta



# ─────────────────────────────────────────────────────────────────────────────
# TOP RESEARCHERS SECTION
# ─────────────────────────────────────────────────────────────────────────────

def _parse_top_researchers(df):
    """
    Parse the raw 'Top Researchers by Impact' sheet (3 cols, header=None style).
    Returns dict with keys: policy, patent, clinical, publications, citations, orcid.
    Each ranked key → list of {rank, name, value}.
    orcid → DataFrame.
    """
    result = {}
    if df.empty:
        return result

    raw = df.copy()
    raw.columns = ['A', 'B', 'C']

    # Identify section headers — col A is text, B and C are NaN
    sections = {}
    current_section = None
    current_rows = []
    in_data = False

    for _, row in raw.iterrows():
        a = str(row['A']).strip() if pd.notna(row['A']) else ''
        b = str(row['B']).strip() if pd.notna(row['B']) else ''
        c = str(row['C']).strip() if pd.notna(row['C']) else ''

        if not a:
            continue

        # Section header: A has text, B and C are blank, and A is not a number
        try:
            int(float(a))
            is_rank = True
        except (ValueError, TypeError):
            is_rank = False

        if not is_rank and b == '' and c == '' and a not in ('Metric', '#'):
            # Save previous section
            if current_section and current_rows:
                sections[current_section] = current_rows
            current_section = a
            current_rows = []
            in_data = False
            continue

        # Skip sub-headers (#, Researcher, ...)
        if a in ('#', 'Metric') or b in ('Researcher',):
            in_data = True
            continue

        if in_data and current_section:
            try:
                rank = int(float(a))
                val_raw = c if c else b
                try:
                    val = int(float(val_raw))
                except (ValueError, TypeError):
                    val = val_raw
                current_rows.append({'Rank': rank, 'Researcher': b, 'Value': val})
            except (ValueError, TypeError):
                # ORCID section has non-rank rows
                if current_section and 'ORCID' in current_section:
                    try:
                        pct = c
                        current_rows.append({'Metric': a, 'Count': b, 'Pct': pct})
                    except Exception:
                        pass

    if current_section and current_rows:
        sections[current_section] = current_rows

    key_map = {
        'Most Policy Citations':        'policy',
        'Most Patent Citations':        'patent',
        'Most Clinical Trial Citations': 'clinical',
        'Most Corporate Citations':     'corporate',
        'Most Publications':            'publications',
        'Most Cited Researchers':       'citations',
        'ORCID Adoption':               'orcid',
    }
    for raw_key, clean_key in key_map.items():
        if raw_key in sections:
            result[clean_key] = sections[raw_key]
    return result


def _build_top_researchers_section(doc, data, meta, tbl):
    """Render the Top Researchers by Impact section."""
    def TBL(t): tbl[0] += 1; return f'Table {tbl[0]}.  {t}'

    tr_raw = data.get('top_researchers', pd.DataFrame())
    if tr_raw.empty:
        return

    parsed = _parse_top_researchers(tr_raw)
    if not parsed:
        return

    _add_heading(doc, 'Top Researchers by Impact', level=2)
    _add_body(doc, (
        f'The following tables identify the highest-performing researchers from the '
        f'{meta.get("faculty", "department")} across key impact categories for the period '
        f'{meta.get("period", "")}. Citation counts are sourced from Dimensions and are '
        f'scoped to the reporting period.'
    ))

    ranked_sections = [
        ('citations',     'Most Cited Researchers',         'Total Citations'),
        ('publications',  'Most Publications',              'Total Publications'),
        ('policy',        'Most Policy Citations',          'Policy Citations'),
        ('patent',        'Most Patent Citations',          'Patent Citations'),
        ('clinical',      'Most Clinical Trial Citations',  'Clinical Trial Citations'),
        ('corporate',     'Most Corporate Citations',       'Corporate Citations'),
    ]

    # If corporate data not in parsed sheet, derive it from Complete Publications
    if not parsed.get('corporate'):
        cp_df = data.get('complete', pd.DataFrame())
        res_df = data.get('researchers', pd.DataFrame())
        if not cp_df.empty and 'corporate_citations' in cp_df.columns and not res_df.empty:
            try:
                cp_exp = cp_df[['researchers','corporate_citations']].copy()
                cp_exp['corporate_citations'] = pd.to_numeric(cp_exp['corporate_citations'], errors='coerce').fillna(0)
                cp_exp = cp_exp.explode('researchers')
                cp_exp['res_id'] = cp_exp['researchers'].apply(
                    lambda x: x.get('id') if isinstance(x, dict) else None)
                cp_exp = cp_exp.dropna(subset=['res_id'])
                corp_by_res = (cp_exp.groupby('res_id')['corporate_citations']
                               .sum().reset_index()
                               .rename(columns={'res_id':'id','corporate_citations':'Corporate Citations'}))
                corp_by_res = corp_by_res[corp_by_res['Corporate Citations'] > 0]
                # Merge researcher names
                name_cols = [c for c in res_df.columns if c in ('first_name','last_name','id')]
                if set(['first_name','last_name','id']).issubset(res_df.columns):
                    names = res_df[['id','first_name','last_name']].copy()
                    names['Researcher'] = names['first_name'].astype(str) + ' ' + names['last_name'].astype(str)
                    corp_by_res = pd.merge(corp_by_res, names[['id','Researcher']], on='id', how='left')
                    corp_by_res = (corp_by_res.dropna(subset=['Researcher'])
                                   .sort_values('Corporate Citations', ascending=False)
                                   .head(10).reset_index(drop=True))
                    corp_by_res['Corporate Citations'] = corp_by_res['Corporate Citations'].astype(int)
                    parsed['corporate'] = [
                        {'Rank': i+1, 'Researcher': row['Researcher'], 'Value': row['Corporate Citations']}
                        for i, row in corp_by_res.iterrows()
                    ]
            except Exception as e:
                print(f'⚠️   Could not derive corporate citations by researcher: {e}')

    for key, heading, value_col in ranked_sections:
        rows = parsed.get(key)
        if not rows:
            continue
        _add_heading(doc, heading, level=3)
        df = pd.DataFrame(rows)
        df = df.rename(columns={'Rank': '#', 'Researcher': 'Researcher', 'Value': value_col})
        _add_table(doc, df,
                   caption=TBL(f'Top researchers by {value_col.lower()}'),
                   col_widths_pct=[6, 66, 28])

    # ORCID adoption table
    orcid_rows = parsed.get('orcid')
    if orcid_rows:
        _add_heading(doc, 'ORCID Adoption', level=3).paragraph_format.keep_with_next = True
        _add_body(doc, (
            'ORCID (Open Researcher and Contributor ID) provides a persistent digital identifier '
            'for researchers. The table below shows ORCID adoption among the researchers included '
            'in this report.'
        )).paragraph_format.keep_with_next = True
        odf = pd.DataFrame(orcid_rows)
        odf.columns = ['Metric', 'Count', '% of Researchers']
        # Remove the "Total researchers" row — always 100%, adds no information
        odf = odf[~odf['Metric'].astype(str).str.strip().str.lower().str.startswith('total res')].reset_index(drop=True)
        _add_table(doc, odf,
                   caption=TBL('ORCID adoption among researchers in this report'),
                   col_widths_pct=[55, 20, 25])


# ─────────────────────────────────────────────────────────────────────────────
# UNIVERSITY COMPARISON SECTION
# ─────────────────────────────────────────────────────────────────────────────

# Short display names for chart labels
_INST_SHORT = {
    'Adelaide University':          'Adelaide',
    'Australian National University':'ANU',
    'Monash University':            'Monash',
    'The University of Sydney':     'Sydney',
    'University of Melbourne':      'Melbourne',
    'University of Queensland':     'UQ',
    'University of Western Australia':'UWA',
}
HOME_INST  = None        # set dynamically at runtime from benchmarking data
HDR_UNSW   = '9B1C1C'   # highlight colour for home institution rows
UNSW_BAR   = '#C0392B'  # accent bar colour for home institution in charts
PEER_BAR   = '#1F497D'  # navy for peer institutions


def _short(name):
    return _INST_SHORT.get(name, name)


def chart_grouped_bar(institutions, metric_values, title, ylabel='',
                      figsize=(8, 4), highlight=None):
    """Single-metric grouped bar, home institution highlighted."""
    if highlight is None:
        highlight = HOME_INST
    fig, ax = plt.subplots(figsize=figsize)
    colors = [UNSW_BAR if i == highlight else PEER_BAR for i in institutions]
    labels = [_short(i) for i in institutions]
    x = range(len(labels))
    bars = ax.bar(x, metric_values, color=colors, edgecolor='white', width=0.6)
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=8, rotation=30, ha='right')
    ax.set_ylabel(ylabel, fontsize=9)
    ax.set_title(title, fontsize=10, fontweight='bold', color=NAVY, pad=8)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v:,.2f}' if isinstance(v, float) and v != int(v) else f'{int(v):,}'))
    for bar, val in zip(bars, metric_values):
        if pd.notna(val) and val > 0:
            lbl = f'{val:,.2f}' if isinstance(val, float) and val != int(val) else f'{int(val):,}'
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height(),
                    lbl, ha='center', va='bottom', fontsize=7)
    ax.set_facecolor('#FAFAFA')
    fig.patch.set_facecolor('white')
    fig.tight_layout()
    return _buf(fig)


def chart_top_cited(df, tier, title, figsize=(8, 4)):
    """Grouped bar: % of pubs in top-cited tier, home institution highlighted."""
    df = df[~df['Institution'].astype(str).str.startswith('ALL')].copy()
    pct_col  = f'Top {tier}% % of Total'
    fcr_col  = f'Top {tier}% FCR (geo mean)'
    df[pct_col] = pd.to_numeric(
        df[pct_col].astype(str).str.replace('%','',regex=False), errors='coerce')
    insts  = df['Institution'].tolist()
    pcts   = df[pct_col].tolist()
    colors = [UNSW_BAR if i == HOME_INST else PEER_BAR for i in insts]
    labels = [_short(i) for i in insts]

    fig, ax = plt.subplots(figsize=figsize)
    x    = range(len(labels))
    bars = ax.bar(x, pcts, color=colors, edgecolor='white', width=0.6)
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=8, rotation=30, ha='right')
    ax.set_ylabel(f'% of publications in top {tier}%', fontsize=9)
    ax.set_title(title, fontsize=10, fontweight='bold', color=NAVY, pad=8)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v:.1f}%'))
    for bar, val in zip(bars, pcts):
        if pd.notna(val):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height(),
                    f'{val:.1f}%', ha='center', va='bottom', fontsize=7)
    ax.set_facecolor('#FAFAFA')
    fig.patch.set_facecolor('white')
    fig.tight_layout()
    return _buf(fig)


def _summary_table_df(df):
    """Clean up a Summary sheet into display-ready DataFrame."""
    col_map = {
        'Institution':              'Institution',
        'Publications':             'Publications',
        'Citations':                'Citations',
        'Citations per paper':      'Cit./Paper',
        'FCR (geo mean)':           'FCR (geo mean)',
        '% Publications with FCR':  '% with FCR',
    }
    available = {c: col_map[c] for c in col_map if c in df.columns}
    df = df[list(available.keys())].copy()
    df = df.rename(columns=available)
    return df.reset_index(drop=True)


def _top_cited_table_df(df):
    """Clean up a Top Cited sheet into display-ready DataFrame."""
    df = df[~df['Institution'].astype(str).str.startswith('ALL')].copy()
    # Support both old names (Top 1% Publications / Top 1% % of Total)
    # and new aggregated names (Publications / Top 1% % of Total)
    rows_out = []
    for _, r in df.iterrows():
        rows_out.append({
            'Institution':  r['Institution'],
            'Top 1% n':   int(r['Publications'] * r.get('Top 1% % of Total', 0) / 100)
                          if 'Publications' in df.columns and pd.notna(r.get('Top 1% % of Total'))
                          else r.get('Top 1% Publications', ''),
            'Top 1% %':   f"{r['Top 1% % of Total']:.1f}%" if pd.notna(r.get('Top 1% % of Total')) else '',
            'Top 5% n':   int(r['Publications'] * r.get('Top 5% % of Total', 0) / 100)
                          if 'Publications' in df.columns and pd.notna(r.get('Top 5% % of Total'))
                          else r.get('Top 5% Publications', ''),
            'Top 5% %':   f"{r['Top 5% % of Total']:.1f}%" if pd.notna(r.get('Top 5% % of Total')) else '',
            'Top 10% n':  int(r['Publications'] * r.get('Top 10% % of Total', 0) / 100)
                          if 'Publications' in df.columns and pd.notna(r.get('Top 10% % of Total'))
                          else r.get('Top 10% Publications', ''),
            'Top 10% %':  f"{r['Top 10% % of Total']:.1f}%" if pd.notna(r.get('Top 10% % of Total')) else '',
        })
    return pd.DataFrame(rows_out).reset_index(drop=True)


def _add_comparison_table(doc, df, caption, unsw_row_color=HDR_UNSW):
    """Like _add_table but highlights the home institution row."""
    n_cols = len(df.columns)
    # Institution names like "University of Western Australia" need at least 35%
    first_w = 35
    rest    = round((100 - first_w) / (n_cols - 1), 1)
    col_pcts = [first_w] + [rest] * (n_cols - 1)
    col_pcts[0] += 100 - sum(col_pcts)

    raw = [int(PAGE_DXA * p / 100) for p in col_pcts]
    raw[0] += PAGE_DXA - sum(raw)

    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(10)
        cp.paragraph_format.space_after  = Pt(3)
        cp.paragraph_format.keep_with_next = True
        _add_run(cp, caption, bold=True, italic=True, size=9.5, color_hex='1F3864')

    tbl = doc.add_table(rows=1 + len(df), cols=n_cols)
    tbl.style = 'Normal Table'

    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tblEl.insert(0, tblPr)
    for ts in tblPr.findall(qn('w:tblStyle')): tblPr.remove(ts)
    for tag in ('w:tblW','w:tblLayout'):
        for el in tblPr.findall(qn(tag)): tblPr.remove(el)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(PAGE_DXA)); tblW.set(qn('w:type'), 'dxa')
    tblPr.insert(0, tblW)
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    for old in tblEl.findall(qn('w:tblGrid')): tblEl.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for w in raw:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGrid.append(gc)
    tblEl.insert(list(tblEl).index(tblPr) + 1, tblGrid)

    def _fmt(cell, ci, text, bold=False, bg='FFFFFF', border_c='C8D8E8',
             align=WD_ALIGN_PARAGRAPH.LEFT, txt_color='1A1A1A'):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for el in tcPr.findall(qn('w:tcW')): tcPr.remove(el)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(raw[ci])); tcW.set(qn('w:type'), 'dxa')
        tcPr.insert(0, tcW)
        for el in tcPr.findall(qn('w:tcBorders')): tcPr.remove(el)
        borders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
            b.set(qn('w:space'),'0'); b.set(qn('w:color'), border_c)
            borders.append(b)
        tcPr.append(borders)
        for el in tcPr.findall(qn('w:shd')): tcPr.remove(el)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), bg.lstrip('#'))
        tcPr.append(shd)
        for el in tcPr.findall(qn('w:tcMar')): tcPr.remove(el)
        mar = OxmlElement('w:tcMar')
        for side, val in (('top',60),('left',100),('bottom',60),('right',100)):
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), str(val)); m.set(qn('w:type'), 'dxa')
            mar.append(m)
        tcPr.append(mar)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        _add_run(p, str(text) if text is not None else '',
                 bold=bold, size=9, color_hex=txt_color)

    # Header
    for ci, col in enumerate(df.columns):
        _fmt(tbl.rows[0].cells[ci], ci, col,
             bold=True, bg=HDR_COLOR, border_c=HDR_COLOR, txt_color='FFFFFF')

    # Data rows
    for ri, (_, row_data) in enumerate(df.iterrows()):
        _keep_row_together(tbl.rows[ri+1])
        inst = str(row_data.iloc[0])
        is_unsw = (HOME_INST is not None and inst == HOME_INST)
        bg = 'FDECEA' if is_unsw else (STRIPE if ri % 2 == 0 else WHITE)
        border_c = 'E8A0A0' if is_unsw else 'C8D8E8'
        for ci, val in enumerate(row_data):
            if pd.isna(val) or val == '' or val is None:
                txt = ''
            elif isinstance(val, float) and val == int(val):
                txt = f'{int(val):,}'
            elif isinstance(val, float):
                txt = f'{val:,.2f}'.rstrip('0').rstrip('.')
            elif isinstance(val, int):
                txt = f'{val:,}'
            else:
                txt = str(val).strip()
            align = WD_ALIGN_PARAGRAPH.RIGHT \
                if ci > 0 and (
                    (isinstance(val, (int, float)) and not pd.isna(val))
                    or txt.lower() == 'n/a'
                ) else WD_ALIGN_PARAGRAPH.LEFT
            txt_color = '7B0000' if is_unsw else '1A1A1A'
            _fmt(tbl.rows[ri+1].cells[ci], ci, txt,
                 bold=is_unsw, bg=bg, border_c=border_c,
                 align=align, txt_color=txt_color)

    _keep_table_on_one_page(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def _build_university_comparison(doc, bench, meta, fig, tbl):
    """Render the full university comparison section into doc."""
    global HOME_INST

    def FIG(t): fig[0] += 1; return f'Fig.{fig[0]}  {t}'
    def TBL(t): tbl[0] += 1; return f'Table {tbl[0]}.  {t}'

    period = meta.get('period', '')

    # Auto-detect home institution: the institution that appears in the
    # Researcher Metrics report's 'Institution' field, matched against
    # the names present in the benchmarking summary sheets.
    home_inst_meta = meta.get('institution', '') or ''
    detected_fors  = bench.get('_fors', [])

    # Collect all institution names from the first available summary sheet
    all_inst_names = []
    for code in detected_fors:
        sdf = bench.get(f'summary_{code}', pd.DataFrame())
        if not sdf.empty and 'Institution' in sdf.columns:
            all_inst_names = sdf['Institution'].dropna().tolist()
            break

    # Match home institution: exact first, then partial
    if home_inst_meta and all_inst_names:
        exact = [n for n in all_inst_names if n == home_inst_meta]
        if exact:
            HOME_INST = exact[0]
        else:
            partial = [n for n in all_inst_names
                       if home_inst_meta.lower() in n.lower()
                       or n.lower() in home_inst_meta.lower()]
            if partial:
                HOME_INST = partial[0]

    # Build the intro description
    home_label = HOME_INST or 'your institution'
    if HOME_INST:
        highlight_note = (f'The row for {HOME_INST} is highlighted in the tables and '
                          f'charts below for easy reference.')
    else:
        highlight_note = ''

    _add_body(doc, (
        f'Comparisons of research performance between universities cannot be carried out at the '
        f'department or school level, as we do not have equivalent lists of researchers from the '
        f'other institutions to compare against. Instead, the comparison is done using publications '
        f'produced by each institution within specific Fields of Research. '
        f'The following comparisons use publications from each institution classified under the '
        f'same Fields of Research (FoR) codes covered in this report, for the period {period}. '
        + (highlight_note if highlight_note else '')
    ))

    # Build section list dynamically from detected FoR codes
    detected_fors = bench.get('_fors', [])
    # ANZSRC 2020 FoR code lookup for common codes
    _FOR_NAMES = {
        '3101': 'Biochemistry and Cell Biology',
        '3102': 'Biophysics',
        '3103': 'Ecology',
        '3104': 'Evolutionary Biology',
        '3105': 'Genetics',
        '3106': 'Industrial Biotechnology',
        '3107': 'Microbiology',
        '3108': 'Plant Biology',
        '3109': 'Zoology',
        '3201': 'Cardiovascular Medicine and Haematology',
        '3202': 'Clinical Sciences',
        '3203': 'Dentistry',
        '3204': 'Epidemiology',
        '3205': 'Medical Biochemistry and Metabolomics',
        '3206': 'Medical Microbiology',
        '3207': 'Neurosciences',
        '3208': 'Oncology and Carcinogenesis',
        '3209': 'Paediatrics',
        '3210': 'Nutrition and Dietetics',
        '3211': 'Ophthalmology and Optometry',
        '3212': 'Orthopaedics',
        '3213': 'Pharmacology and Pharmaceutical Sciences',
        '3214': 'Reproductive Medicine',
        '3215': 'Geriatrics',
        '3301': 'Architecture',
        '3302': 'Building',
        '3303': 'Urban and Regional Planning',
        '3401': 'Aerospace Engineering',
        '3402': 'Automotive Engineering',
        '3403': 'Chemical Engineering',
        '3404': 'Civil Engineering',
        '3405': 'Electrical and Electronic Engineering',
        '3406': 'Environmental Engineering',
        '3407': 'Food Sciences',
        '3408': 'Geomatic Engineering',
        '3409': 'Manufacturing Engineering',
        '3410': 'Maritime Engineering',
        '3411': 'Mechanical Engineering',
        '3412': 'Nanotechnology',
        '3501': 'Accounting, Auditing and Accountability',
        '3502': 'Banking, Finance and Investment',
        '3503': 'Business Systems in Context',
        '3504': 'Commercial Services',
        '3505': 'Human Resources and Industrial Relations',
        '3506': 'Marketing',
        '3507': 'Strategy, Management and Organisational Behaviour',
        '3601': 'Art History, Theory and Criticism',
        '3602': 'Creative and Professional Writing',
        '3603': 'Music',
        '3604': 'Performing Arts',
        '3605': 'Screen and Digital Media',
        '3606': 'Visual Arts',
        '4001': 'Anthropology',
        '4002': 'Archaeology',
        '4003': 'Criminology',
        '4004': 'Development Studies',
        '4005': 'Education Systems',
        '4006': 'Gender Studies',
        '4007': 'Human Geography',
        '4008': 'Political Science',
        '4009': 'Social Work',
        '4010': 'Sociology',
        '4101': 'Climate Change Impacts and Adaptation',
        '4102': 'Ecological Applications',
        '4103': 'Environmental Biotechnology',
        '4104': 'Environmental Management',
        '4105': 'Pollution and Contamination',
        '4106': 'Soil Sciences',
        '4601': 'Applied Computing',
        '4602': 'Artificial Intelligence',
        '4603': 'Computer Vision and Multimedia Computation',
        '4604': 'Cybersecurity and Privacy',
        '4605': 'Data Management and Data Science',
        '4606': 'Distributed Computing and Systems Software',
        '4607': 'Graphics, Augmented Reality and Games',
        '4608': 'Human-Centred Computing',
        '4609': 'Information Systems',
        '4610': 'Library and Information Studies',
        '4611': 'Machine Learning',
        '4612': 'Software Engineering',
        '5101': 'Astronomical Sciences',
        '5102': 'Atomic, Molecular and Optical Physics',
        '5103': 'Classical Physics',
        '5104': 'Condensed Matter Physics',
        '5105': 'Medical and Biological Physics',
        '5106': 'Nuclear and Plasma Physics',
        '5107': 'Particle and High Energy Physics',
        '5108': 'Quantum Physics',
        '5201': 'Applied Mathematics',
        '5202': 'Ecological Modelling',
        '5203': 'Mathematical Physics',
        '5204': 'Pure Mathematics',
        '5205': 'Statistics',
        '4001': 'Geochemistry',
        '3702': 'Geophysics',
        '3703': 'Hydrography',
        '3704': 'Oceanography',
        '3705': 'Geological Sciences',
    }
    for_sections = [('All Fields of Research', 'summary', 'top_cited', 'all FoR codes')]
    for code in detected_fors:
        name = _FOR_NAMES.get(str(code), '')
        raw_label = f'FoR {code} – {name}' if name else f'FoR {code}'
        for_sections.append((raw_label, f'summary_{code}', f'top_cited_{code}', f'FoR {code}'))

    for label, s_key, tc_key, for_desc in for_sections:
        s_df  = bench.get(s_key,  pd.DataFrame())
        tc_df = bench.get(tc_key, pd.DataFrame())
        if s_df.empty and tc_df.empty:
            continue

        _add_heading(doc, label, level=3)

        # ── Summary table ──────────────────────────────────────────────────
        if not s_df.empty:
            tbl_df = _summary_table_df(s_df)
            _add_comparison_table(
                doc, tbl_df,
                caption=TBL(f'Research output summary — {label}')
            )

        # ── Chart: publications ────────────────────────────────────────────
        if not s_df.empty:
            s_plot = s_df[~s_df['Institution'].astype(str).str.startswith('ALL')].copy()
            insts  = s_plot['Institution'].tolist()

            buf = chart_grouped_bar(
                insts,
                pd.to_numeric(s_plot['Publications'], errors='coerce').tolist(),
                f'Total Publications — {label}',
                ylabel='Publications'
            )
            _insert_image(doc, buf, 6.2,
                FIG(f'Total publications by institution — {label}'))

            buf = chart_grouped_bar(
                insts,
                pd.to_numeric(s_plot['FCR (geo mean)'], errors='coerce').tolist(),
                f'FCR (Geometric Mean) — {label}',
                ylabel='FCR (geo mean)'
            )
            _insert_image(doc, buf, 6.2,
                FIG(f'Field Citation Ratio (geometric mean) by institution — {label}'))

            buf = chart_grouped_bar(
                insts,
                pd.to_numeric(s_plot['Citations per paper'], errors='coerce').tolist(),
                f'Citations per Paper — {label}',
                ylabel='Citations per paper'
            )
            _insert_image(doc, buf, 6.2,
                FIG(f'Citations per paper by institution — {label}'))

        # ── Top cited table ────────────────────────────────────────────────
        if not tc_df.empty:
            tc_tbl = _top_cited_table_df(tc_df)
            _add_comparison_table(
                doc, tc_tbl,
                caption=TBL(f'Top cited publications — {label}')
            )

            # Top cited charts — one per tier
            for tier in (1, 5, 10):
                pct_col = f'Top {tier}% % of Total'
                tc_plot = tc_df[~tc_df['Institution'].astype(str).str.startswith('ALL')].copy()
                tc_plot[pct_col] = pd.to_numeric(
                    tc_plot[pct_col].astype(str).str.replace('%','',regex=False),
                    errors='coerce')
                if tc_plot[pct_col].notna().any():
                    buf = chart_top_cited(
                        tc_df, tier,
                        f'Top {tier}% Most-Cited Publications — {label}'
                    )
                    _insert_image(doc, buf, 6.2,
                        FIG(f'% of publications in top {tier}% most-cited globally — {label}'))


# ─────────────────────────────────────────────────────────────────────────────
# SECTION BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

def section_cover(doc, meta):
    institution = meta.get('institution', '') or ''
    # Strip any trailing NaN or placeholder text
    if institution.lower() in ('nan', 'none', ''):
        institution = ''

    if institution:
        p = doc.add_paragraph()
        _add_run(p, institution, bold=True, size=18, color_hex='156082')

    p2 = doc.add_paragraph()
    faculty = meta.get('faculty', '')
    _add_run(p2, f'Report for the {faculty}' if faculty else 'Research Report',
             bold=True, size=14)

    p3 = doc.add_paragraph()
    _add_run(p3, meta.get('period', ''), size=13)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(24)
    _add_run(p4, f'Prepared on {meta.get("date", "")}',
             size=10, color_hex='156082')


def section_part1(doc, data, meta):
    # Counter shared across sections — pass as list so it's mutable
    # (caller manages fig_n and tbl_n)
    pass


# ─────────────────────────────────────────────────────────────────────────────
# FULL REPORT BUILDER  (manages all figure/table numbering)
# ─────────────────────────────────────────────────────────────────────────────

def _add_contents_page(doc, has_datasets=False, has_corporate=False, has_benchmarking=False):
    """Build a hyperlinked Table of Contents matching the departmental report structure."""

    def _toc_entry(text, level=1, bookmark=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2 if level > 1 else 6)
        p.paragraph_format.space_after  = Pt(2 if level > 1 else 4)
        p.paragraph_format.keep_with_next = False
        if level > 1:
            p.paragraph_format.left_indent = Pt(22)
        bm = bookmark if bookmark else _make_bookmark_id(text)
        if level == 1:
            _add_internal_hyperlink(p, text, bm, bold=True,  size=11, color_hex='1F3864')
        else:
            _add_internal_hyperlink(p, text, bm, bold=False, size=10, color_hex='434343')

    def _divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), 'C8D8E8')
        pBdr.append(bot); pPr.append(pBdr)

    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(16)
    _add_run(p_title, 'Contents', bold=True, size=16, color_hex='1F3864')

    # ── Part 1 ────────────────────────────────────────────────────────────────
    _toc_entry('Part 1: Research Outputs data', level=1)
    _toc_entry('Total publications by year',          level=2)
    _toc_entry('Total publications by type',          level=2)
    _toc_entry('Aggregated Publication Indicators',   level=2)
    _toc_entry('Top 10 Journals by Publication Count',level=2)
    _toc_entry('Funding',                             level=2)
    if has_datasets:
        _toc_entry('Datasets',                        level=2)
    if has_corporate:
        _toc_entry('Corporate Citations',             level=2)
    _divider()

    # ── Part 2 ────────────────────────────────────────────────────────────────
    _toc_entry('Part 2: Metrics from Dimensions', level=1)
    _toc_entry('Fields of Research',                          level=2)
    _toc_entry('Field Citation Ratio – FCR (Geometric mean)', level=2)
    _toc_entry('Authorship collaborations',                   level=2)
    _toc_entry('International Collaboration by Country',      level=2)
    _toc_entry('Industry Co-authorship',                      level=2)
    _toc_entry('Government Co-authorship',                    level=2)
    _toc_entry('Healthcare Co-authorship',                    level=2)
    _toc_entry('Facility Co-authorship',                      level=2)
    _toc_entry('Nonprofit Co-authorship',                     level=2)
    _toc_entry('Top Researchers by Impact',                   level=2,
               bookmark=_make_bookmark_id('Top Researchers by Impact'))
    _divider()

    # ── Part 3 ────────────────────────────────────────────────────────────────
    _toc_entry('Part 3: Pathways to Impact', level=1)
    _toc_entry('Summary of Impact',              level=2)
    _toc_entry('Policy Citations',               level=2)
    _toc_entry('Patent Citations',               level=2)
    _toc_entry('Clinical Trial Citations',       level=2)
    _toc_entry('Most Alternative Research Attention', level=2)
    _divider()

    # ── Part 4 ────────────────────────────────────────────────────────────────
    if has_benchmarking:
        _toc_entry('Part 4: University Comparison by Field of Research', level=1)
        _divider()

    # ── Appendix ──────────────────────────────────────────────────────────────
    _toc_entry('Appendix: Methodology & Report Information', level=1)

    doc.add_page_break()



def build_report(doc, data, meta):
    cp = data['complete']
    ps = data['pub_summary']
    yr_c = [c for c in ps.columns if re.match(r'^\d{4}$', str(c))]

    fig = [0]   # mutable counter
    tbl = [0]

    # ── Contents page ──────────────────────────────────────────────────────────
    has_datasets    = not data.get('datasets', pd.DataFrame()).empty
    has_corporate   = ('corporate_citations' in cp.columns and
                       pd.to_numeric(cp.get('corporate_citations', pd.Series()), errors='coerce').gt(0).any())
    has_benchmarking = bool(data.get('bench', {}))
    _add_contents_page(doc,
                       has_datasets=has_datasets,
                       has_corporate=has_corporate,
                       has_benchmarking=has_benchmarking)

    def FIG(caption_text):
        fig[0] += 1
        return f'Fig.{fig[0]}  {caption_text}'

    def TBL(caption_text):
        tbl[0] += 1
        return f'Table {tbl[0]}.  {caption_text}'

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1
    # ══════════════════════════════════════════════════════════════════════════
    _add_heading(doc, 'Part 1: Research Outputs data', level=1)
    _add_body(doc, (
        'This data has been generated from the Dimensions publications dataset, '
        'using a list of researchers from the research management system.'
    ))

    # ── Fig.1 – Publications by Year ─────────────────────────────────────────
    _add_heading(doc, 'Total publications by year', level=3)
    pub_row = ps[ps['Metric'] == 'Total Publications']
    if not pub_row.empty and yr_c:
        values = [float(pub_row[c].values[0]) for c in yr_c]
        buf = chart_col(yr_c, values, 'Publications by Year',
                        xlabel='Year', ylabel='No. Publications', figsize=(7, 3.5))
        _insert_image(doc, buf, 6.0, FIG('Total number of publications by year'))

    # ── Fig.2 – Publications by Type ─────────────────────────────────────────
    _add_heading(doc, 'Total publications by type', level=3)
    if 'type' in cp.columns:
        tc = cp['type'].value_counts()
        buf = chart_col(tc.index.tolist(), tc.values.tolist(),
                        'Publications by Type',
                        xlabel='Publication Type', ylabel='No. Publications',
                        figsize=(7, 3.5))
        _insert_image(doc, buf, 6.0, FIG('Total number of publications by type'))

    # ── Table 1 – Aggregated Publication Indicators ───────────────────────────
    _add_heading(doc, 'Aggregated Publication Indicators', level=3)
    imp = data['impact'][['Metric', 'Count']].copy()
    imp['Count'] = pd.to_numeric(imp['Count'], errors='coerce')

    # Replace Mean Altmetric Score with % Publications with Altmetric Attention
    mask_mean = imp['Metric'] == 'Mean Altmetric Score'
    if mask_mean.any():
        total_pubs = imp.loc[imp['Metric'] == 'Total Publications', 'Count'].values
        with_alt   = imp.loc[imp['Metric'] == 'Publications with Altmetric Score', 'Count'].values
        if len(total_pubs) and len(with_alt) and total_pubs[0] > 0:
            pct = round(with_alt[0] / total_pubs[0] * 100, 1)
            imp.loc[mask_mean, 'Metric'] = '% Publications with Altmetric Attention'
            imp.loc[mask_mean, 'Count']  = pct

    pct_metric = '% Publications with Altmetric Attention'
    def _fmt_count(row):
        v = row['Count']
        if pd.isna(v):
            return ''
        if row['Metric'] == pct_metric:
            return f'{v:.1f}%'
        if v == int(v):
            return f'{int(v):,}'
        return f'{v:,.3f}'

    imp['Count'] = imp.apply(_fmt_count, axis=1)
    imp.columns = ['Metric', 'Value']
    _add_table(doc, imp,
               caption=TBL('Aggregated Publication Indicators'),
               col_widths_pct=[78, 22])

    # ── Table 2 – Top 10 Journals ─────────────────────────────────────────────
    _add_heading(doc, 'Top 10 Journals by Publication Count', level=3)
    jraw = data['journals']
    top10_col = 'Top 10 Journals by Publication Count'
    if top10_col in jraw.columns:
        start = jraw.columns.get_loc(top10_col)
        jdf = jraw.iloc[1:, start:start+4].copy()
        jdf.columns = ['Journal', 'Pubs', 'SNIP', 'SJR']
        jdf = jdf.dropna(subset=['Journal']).reset_index(drop=True)
        jdf['Pubs'] = pd.to_numeric(jdf['Pubs'], errors='coerce').apply(
            lambda v: int(v) if pd.notna(v) else '')
        for col in ('SNIP','SJR'):
            jdf[col] = pd.to_numeric(jdf[col], errors='coerce').apply(
                lambda v: round(v, 3) if pd.notna(v) else 'n/a')
    else:
        jdf = jraw[['journal.title','Count','SNIP','SJR']].copy()
        jdf.columns = ['Journal','Pubs','SNIP','SJR']
        jdf = jdf.sort_values('Pubs', ascending=False).head(10).reset_index(drop=True)
    _add_table(doc, jdf,
               caption=TBL('Top 10 Journals by publication output with SJR and SNIP (where available).'),
               col_widths_pct=[52, 14, 17, 17])

    # ── Table 3 – Top 10 Funders ──────────────────────────────────────────────
    _add_heading(doc, 'Funding', level=3)
    fdf = data['funders'][['Funder','Publications']].copy() \
          if 'Publications' in data['funders'].columns \
          else data['funders'][['Funder','Count']].rename(columns={'Count':'Publications'})
    fdf = fdf.dropna(subset=['Publications']).sort_values('Publications', ascending=False).head(10).reset_index(drop=True)
    _add_table(doc, fdf,
               caption=TBL('Top 10 funders by number of supported publications, highlighting the principal sources of research support.'),
               col_widths_pct=[76, 24])

    # ── Datasets ──────────────────────────────────────────────────────────────
    ds_df = data.get('datasets', pd.DataFrame())
    if not ds_df.empty:
        _add_heading(doc, 'Datasets', level=3)
        faculty_str = meta.get('faculty', 'this department')
        _add_body(doc, (
            f'The following datasets were produced by researchers in {faculty_str} '
            f'during {meta.get("period", "the reporting period")} and are recorded in Dimensions. '
            f'Datasets provide a direct measure of research outputs beyond traditional publications.'
        ))
        try:
            if 'year' in ds_df.columns:
                ds_df['year'] = pd.to_numeric(ds_df['year'], errors='coerce')
                yr_ds = (ds_df.groupby('year').size().reset_index(name='Count')
                         .dropna(subset=['year']))
                yr_ds['year'] = yr_ds['year'].astype(int)
                if len(yr_ds) > 1:
                    buf = chart_col(yr_ds['year'].astype(str).tolist(),
                                    yr_ds['Count'].tolist(),
                                    'Datasets by Year',
                                    xlabel='Year', ylabel='No. Datasets', figsize=(6, 3))
                    _insert_image(doc, buf, 5.5, FIG('Total number of datasets by year'))
            title_col = next((c for c in ds_df.columns if 'title' in str(c).lower()), None)
            year_col  = next((c for c in ds_df.columns if 'year'  in str(c).lower()), None)
            doi_col   = next((c for c in ds_df.columns if 'doi'   in str(c).lower()), None)
            if title_col:
                sort_col = year_col if year_col else title_col
                ds_show = ds_df.sort_values(sort_col, ascending=False).head(20).reset_index(drop=True)
                rows_ds = []
                for _, r in ds_show.iterrows():
                    doi = str(r[doi_col]) if doi_col and pd.notna(r[doi_col]) else ''
                    url = f'https://doi.org/{doi}' if doi and doi.lower() not in ('nan', '') else ''
                    yr  = int(r[year_col]) if year_col and pd.notna(r[year_col]) else ''
                    rows_ds.append({'Title': str(r[title_col]), 'Year': yr, '_url': url})
                _add_hyperlink_table(
                    doc, rows_ds,
                    col_headers=['Title', 'Year'],
                    col_widths_pct=[82, 18],
                    url_col_idx=0,
                    caption=TBL('Datasets')
                )
        except Exception as e:
            print(f'⚠️   Could not build Datasets table: {e}')

    # ── Table 3a – Corporate Citations ────────────────────────────────────────
    cc_raw = data.get('corporate_citations', pd.DataFrame())
    if not cc_raw.empty:
        _add_heading(doc, 'Corporate Citations', level=3)
        # Sheet layout: col 0 = Organisation, col 1 = Corporate Citations, col 2 = Publications Cited
        # Row 0 is the title note; row 1 is the column header; data starts at row 2.
        # There are TWO stacked tables — we only want the first (≥2 citations, recommended).
        # Stop at the first blank row after the data starts.
        try:
            cc_raw.columns = range(len(cc_raw.columns))
            # Find the end of the first table (first blank in col 0 after row 2)
            end_row = None
            for i in range(2, len(cc_raw)):
                if pd.isna(cc_raw.iloc[i, 0]) or str(cc_raw.iloc[i, 0]).strip() == '':
                    end_row = i
                    break
            cc_df = cc_raw.iloc[2:end_row, [0, 1, 2]].copy() if end_row else cc_raw.iloc[2:, [0, 1, 2]].copy()
            cc_df.columns = ['Organisation', 'Corporate Citations', 'Publications Cited']
            cc_df = cc_df[cc_df['Organisation'].notna()].copy()
            cc_df = cc_df[cc_df['Organisation'].astype(str).str.strip().str.lower() != 'organisation'].copy()
            cc_df['Corporate Citations'] = pd.to_numeric(cc_df['Corporate Citations'], errors='coerce')
            cc_df['Publications Cited']  = pd.to_numeric(cc_df['Publications Cited'],  errors='coerce')
            cc_df = cc_df.dropna(subset=['Corporate Citations'])
            cc_df = cc_df.sort_values('Corporate Citations', ascending=False).head(20).reset_index(drop=True)
            cc_df['Corporate Citations'] = cc_df['Corporate Citations'].apply(lambda v: f'{int(v):,}')
            cc_df['Publications Cited']  = cc_df['Publications Cited'].apply(
                lambda v: f'{int(v):,}' if pd.notna(v) else '')
            _add_table(doc, cc_df,
                       caption=TBL('Top 20 organisations by corporate citations '
                                   '(organisations with \u22652 corporate citations).'),
                       col_widths_pct=[62, 19, 19])
        except Exception as e:
            print(f'⚠️   Could not build Corporate Citations table: {e}')


    # ══════════════════════════════════════════════════════════════════════════
    _page_break(doc)
    _add_heading(doc, 'Part 2: Metrics from Dimensions', level=1)

    # ── Fig.3 & 4 – Fields of Research ───────────────────────────────────────
    _add_heading(doc, 'Fields of Research', level=3)
    for_df = data['for_'][['Field','Count']].copy()
    for_df['Count'] = pd.to_numeric(for_df['Count'], errors='coerce')
    for_df = for_df.dropna(subset=['Count'])
    for_df['nd'] = for_df['Field'].apply(
        lambda x: len(re.match(r'^(\d+)', str(x)).group(1))
        if re.match(r'^(\d+)', str(x)) else 0)

    df_2d = for_df[for_df['nd']==2].sort_values('Count',ascending=False).head(10).reset_index(drop=True)
    df_4d = for_df[for_df['nd']==4].sort_values('Count',ascending=False).head(10).reset_index(drop=True)

    if not df_2d.empty:
        buf = chart_hbar(df_2d['Field'].tolist()[::-1], df_2d['Count'].tolist()[::-1],
                         'Top 10 2-Digit Fields of Research',
                         xlabel='No. Publications', figsize=(7, 4))
        _insert_image(doc, buf, 6.0,
                      FIG('Top 10 subject areas by 2-Digit Field of Research (ANZSRC 2020)'))

    if not df_4d.empty:
        buf = chart_hbar(df_4d['Field'].tolist()[::-1], df_4d['Count'].tolist()[::-1],
                         'Top 10 4-Digit Fields of Research',
                         xlabel='No. Publications', figsize=(7, 4.5))
        _insert_image(doc, buf, 6.0,
                      FIG('Top 10 subfield area coverage by 4-Digit Field of Research (ANZSRC 2020)'))

    # ── Table 4 – FCR by Year ─────────────────────────────────────────────────
    _add_heading(doc, 'Field Citation Ratio \u2013 FCR (Geometric mean)', level=3).paragraph_format.keep_with_next = True
    _add_body(doc, 'The yearly and overall Geometric mean field citation ratio (FCR) for publications.').paragraph_format.keep_with_next = True
    fcr_row = ps[ps['Metric'].str.contains('Field Citation|FCR', na=False, regex=True)]
    if not fcr_row.empty and yr_c:
        def _fcr(col):
            try:
                v = fcr_row[col].values[0]
                return round(float(v), 3) if pd.notna(v) else 'n/a'
            except Exception:
                return 'n/a'
        fcr_df = pd.DataFrame({
            'Period': ['Overall'] + list(yr_c),
            'FCR (Geometric Mean)': [_fcr('Overall')] + [_fcr(c) for c in yr_c]
        })
        _add_table(doc, fcr_df,
                   caption=TBL('Field citation ratio \u2013 FCR (geometric mean) for publications'),
                   col_widths_pct=[30, 70])

    # ── Fig.5 – Collaboration by type ─────────────────────────────────────────
    _add_heading(doc, 'Authorship collaborations', level=3)
    _add_body(doc, ('Publications collaborations by type of co-authorship (international, national, '
                    'institutional and single-authored) are shown in Fig\u00a05 and by organisation type in Fig\u00a06.'))
    col_df = data['collab'].copy()
    ORDER = ['International Collaboration','Only national collaboration',
             'Only institutional collaboration','Single authorship (no collaboration)']
    SHORT = ['International','National','Institutional','Single']
    if 'Collaboration_Type' in col_df.columns and 'Count' in col_df.columns:
        col_df = col_df.set_index('Collaboration_Type').reindex(ORDER).fillna(0).reset_index()
        buf = chart_col(SHORT, col_df['Count'].tolist(),
                        'Publications by Collaboration Type',
                        ylabel='No. Publications', figsize=(6, 3.5))
        _insert_image(doc, buf, 5.5, FIG('Publications by collaboration type'))

    # ── Fig.6 – Org type ─────────────────────────────────────────────────────
    org_df = data['org_collab'].copy()
    if 'Organisation Type' in org_df.columns and 'Publications' in org_df.columns:
        org_df = (org_df.dropna(subset=['Organisation Type'])
                        .query('`Organisation Type`.str.strip() != ""', engine='python')
                        [['Organisation Type','Publications']]
                        .sort_values('Publications', ascending=False).head(8))
        org_df['Publications'] = pd.to_numeric(org_df['Publications'], errors='coerce').fillna(0)
        buf = chart_col(org_df['Organisation Type'].tolist(),
                        org_df['Publications'].tolist(),
                        'Publications by Organisation Type',
                        ylabel='No. Publications', figsize=(6.5, 3.5))
        _insert_image(doc, buf, 5.5, FIG('Publications by organisation type'))

    # ── Fig.7 – Country Collaboration (if tab exists) ─────────────────────────
    cc = data['country_collab']
    # Normalise column name — new export uses 'Country (excl. Australia)'
    if not cc.empty and 'Country (excl. Australia)' in cc.columns and 'Country' not in cc.columns:
        cc = cc.rename(columns={'Country (excl. Australia)': 'Country'})
    if not cc.empty and 'Country' in cc.columns and 'Publications' in cc.columns:
        _add_heading(doc, 'International Collaboration by Country', level=3)
        cc_top = cc.dropna(subset=['Country']).head(20).copy()
        cc_top['Publications'] = pd.to_numeric(cc_top['Publications'], errors='coerce').fillna(0)
        buf = chart_hbar(cc_top['Country'].tolist()[::-1], cc_top['Publications'].tolist()[::-1],
                         'Top 20 Collaborating Countries',
                         xlabel='No. Publications', figsize=(7, 6))
        _insert_image(doc, buf, 6.0, FIG('Top 20 collaborating countries by number of publications'))

    # ── Industry Co-authorship ────────────────────────────────────────────────
    _add_heading(doc, 'Industry Co-authorship', level=3)
    _add_body(doc, (
        'The following table shows organisations that have co-authored publications with '
        'researchers in this report, where the co-authoring organisation is classified as '
        'a Company in Dimensions. This reflects active research partnerships rather than '
        'citations, and is a direct measure of industry engagement in research activity.'
    ))
    ind_df = data.get('industry_coauth', pd.DataFrame())
    if not ind_df.empty:
        try:
            ind_df = ind_df.copy()
            ind_df.columns = [str(c) for c in ind_df.columns]
            org_col  = ind_df.columns[0]
            pubs_col = ind_df.columns[1]
            cit_col  = ind_df.columns[2]
            ind_df[pubs_col] = pd.to_numeric(ind_df[pubs_col], errors='coerce')
            ind_df = ind_df.dropna(subset=[pubs_col])
            ind_df = ind_df.sort_values(pubs_col, ascending=False).head(20).reset_index(drop=True)
            ind_df[pubs_col] = ind_df[pubs_col].apply(lambda v: f'{int(v):,}')
            ind_df[cit_col]  = pd.to_numeric(ind_df[cit_col], errors='coerce').apply(
                lambda v: f'{int(v):,}' if pd.notna(v) else '')
            _add_table(doc, ind_df[[org_col, pubs_col, cit_col]],
                       caption=TBL('Top 20 industry co-authoring organisations (Company-type, ranked by co-authored publications)'),
                       col_widths_pct=[62, 19, 19])
            # Top 10 co-authored publications with Dimensions links
            comp = data.get('complete', pd.DataFrame())
            if not comp.empty and 'research_org_types' in comp.columns:
                def _has_co(val):
                    if isinstance(val, list):
                        return any('company' in str(t).lower() for t in val)
                    return 'company' in str(val).lower()
                top_ind_pubs = comp[comp['research_org_types'].apply(_has_co)].copy()
                top_ind_pubs['times_cited'] = pd.to_numeric(
                    top_ind_pubs.get('times_cited', 0), errors='coerce').fillna(0)
                top_ind_pubs = (top_ind_pubs[['id','title','doi','year','times_cited']]
                                .sort_values('times_cited', ascending=False)
                                .head(10).reset_index(drop=True))
                if not top_ind_pubs.empty:
                    _add_body(doc, 'The top 10 most cited publications co-authored with industry partners are shown below.').paragraph_format.keep_with_next = True
                    rows_ind = []
                    for _, r in top_ind_pubs.iterrows():
                        pub_id = str(r['id']) if pd.notna(r.get('id', '')) else ''
                        url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
                        rows_ind.append({
                            'Publication Title': str(r['title']),
                            'Year': int(r['year']) if pd.notna(r['year']) else '',
                            'Citations': int(r['times_cited']),
                            '_url': url
                        })
                    _add_hyperlink_table(
                        doc, rows_ind,
                        col_headers=['Publication Title', 'Year', 'Citations'],
                        col_widths_pct=[75, 10, 15],
                        url_col_idx=0,
                        caption=TBL('Top 10 publications co-authored with industry (Company-type organisations), ranked by citations')
                    )
        except Exception as e:
            print(f'⚠️   Could not build Industry Co-authorship table: {e}')

    # ── Government Co-authorship ─────────────────────────────────────────────
    _add_heading(doc, 'Government Co-authorship', level=3)
    _add_body(doc, (
        'The following table shows the top publications co-authored with researchers in this '
        'report where at least one co-authoring organisation is classified as Government in '
        'Dimensions. This reflects active research partnerships and is a direct measure of '
        'engagement with government organisations.'
    ))
    try:
        if not cp.empty and 'research_org_types' in cp.columns:
            def _has_govt(val):
                if isinstance(val, list):
                    return any('government' in str(t).lower() for t in val)
                return 'government' in str(val).lower()
            govt_pubs = cp[cp['research_org_types'].apply(_has_govt)].copy()
            if not govt_pubs.empty:
                govt_pubs['times_cited'] = pd.to_numeric(
                    govt_pubs.get('times_cited', 0), errors='coerce').fillna(0)
                # Year chart
                yr_gov = (govt_pubs.groupby('year').size().reset_index(name='Count')
                          .dropna(subset=['year']))
                yr_gov['year'] = yr_gov['year'].astype(int)
                if len(yr_gov) > 1:
                    buf = chart_col(yr_gov['year'].astype(str).tolist(),
                                    yr_gov['Count'].tolist(),
                                    'Government Co-authored Publications by Year',
                                    xlabel='Year', ylabel='No. Publications', figsize=(6, 3))
                    _insert_image(doc, buf, 5.5,
                                  FIG('Government co-authored publications by year'))
                top_gov = (govt_pubs[['id','title','year','times_cited']]
                           .sort_values('times_cited', ascending=False)
                           .head(10).reset_index(drop=True))
                rows_gov = []
                for _, r in top_gov.iterrows():
                    pub_id = str(r['id']) if pd.notna(r.get('id', '')) else ''
                    url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
                    rows_gov.append({
                        'Publication Title': str(r['title']),
                        'Year': int(r['year']) if pd.notna(r['year']) else '',
                        'Citations': int(r['times_cited']),
                        '_url': url
                    })
                _add_hyperlink_table(
                    doc, rows_gov,
                    col_headers=['Publication Title', 'Year', 'Citations'],
                    col_widths_pct=[75, 10, 15],
                    url_col_idx=0,
                    caption=TBL('Top publications co-authored with government organisations, ranked by citations')
                )
            else:
                _add_body(doc, 'No government co-authored publications found.', italic=True)
        else:
            _add_body(doc, '(Organisation type data not available.)', italic=True)
    except Exception as e:
        print(f'⚠️   Could not build Government Co-authorship section: {e}')

    # ── Healthcare Co-authorship ──────────────────────────────────────────────
    _add_heading(doc, 'Healthcare Co-authorship', level=3)
    _add_body(doc, (
        'The following table shows the top publications co-authored with researchers in this '
        'report where at least one co-authoring organisation is classified as Healthcare or '
        'Medical in Dimensions. This reflects active research partnerships with healthcare '
        'and medical organisations.'
    ))
    try:
        if not cp.empty and 'research_org_types' in cp.columns:
            def _has_health(val):
                if isinstance(val, list):
                    return any(kw in str(t).lower() for kw in ['healthcare', 'medical'] for t in val)
                return any(kw in str(val).lower() for kw in ['healthcare', 'medical'])
            health_pubs = cp[cp['research_org_types'].apply(_has_health)].copy()
            if not health_pubs.empty:
                health_pubs['times_cited'] = pd.to_numeric(
                    health_pubs.get('times_cited', 0), errors='coerce').fillna(0)
                yr_hlt = (health_pubs.groupby('year').size().reset_index(name='Count')
                          .dropna(subset=['year']))
                yr_hlt['year'] = yr_hlt['year'].astype(int)
                if len(yr_hlt) > 1:
                    buf = chart_col(yr_hlt['year'].astype(str).tolist(),
                                    yr_hlt['Count'].tolist(),
                                    'Healthcare Co-authored Publications by Year',
                                    xlabel='Year', ylabel='No. Publications', figsize=(6, 3))
                    _insert_image(doc, buf, 5.5,
                                  FIG('Healthcare co-authored publications by year'))
                top_hlt = (health_pubs[['id','title','year','times_cited']]
                           .sort_values('times_cited', ascending=False)
                           .head(10).reset_index(drop=True))
                rows_hlt = []
                for _, r in top_hlt.iterrows():
                    pub_id = str(r['id']) if pd.notna(r.get('id', '')) else ''
                    url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
                    rows_hlt.append({
                        'Publication Title': str(r['title']),
                        'Year': int(r['year']) if pd.notna(r['year']) else '',
                        'Citations': int(r['times_cited']),
                        '_url': url
                    })
                _add_hyperlink_table(
                    doc, rows_hlt,
                    col_headers=['Publication Title', 'Year', 'Citations'],
                    col_widths_pct=[75, 10, 15],
                    url_col_idx=0,
                    caption=TBL('Top publications co-authored with healthcare and medical organisations, ranked by citations')
                )
            else:
                _add_body(doc, 'No healthcare co-authored publications found.', italic=True)
        else:
            _add_body(doc, '(Organisation type data not available.)', italic=True)
    except Exception as e:
        print(f'⚠️   Could not build Healthcare Co-authorship section: {e}')

    # ── Facility Co-authorship ───────────────────────────────────────────────
    _add_heading(doc, 'Facility Co-authorship', level=3)
    _add_body(doc, (
        'The following table shows publications co-authored with researchers in this '
        'report where at least one co-authoring organisation is classified as a Facility '
        'in Dimensions. This includes research facilities, national laboratories, and '
        'large-scale research infrastructure.'
    ))
    try:
        if not cp.empty and 'research_org_types' in cp.columns:
            def _has_facility(val):
                if isinstance(val, list):
                    return any('facility' in str(t).lower() for t in val)
                return 'facility' in str(val).lower()
            facility_pubs = cp[cp['research_org_types'].apply(_has_facility)].copy()
            if not facility_pubs.empty:
                facility_pubs['times_cited'] = pd.to_numeric(
                    facility_pubs.get('times_cited', 0), errors='coerce').fillna(0)
                yr_fac = (facility_pubs.groupby('year').size().reset_index(name='Count')
                          .dropna(subset=['year']))
                yr_fac['year'] = yr_fac['year'].astype(int)
                if len(yr_fac) > 1:
                    buf = chart_col(yr_fac['year'].astype(str).tolist(),
                                    yr_fac['Count'].tolist(),
                                    'Facility Co-authored Publications by Year',
                                    xlabel='Year', ylabel='No. Publications', figsize=(6, 3))
                    _insert_image(doc, buf, 5.5,
                                  FIG('Facility co-authored publications by year'))
                top_fac = (facility_pubs[['id', 'title', 'year', 'times_cited']]
                           .sort_values('times_cited', ascending=False)
                           .head(10).reset_index(drop=True))
                rows_fac = []
                for _, r in top_fac.iterrows():
                    pub_id = str(r['id']) if pd.notna(r.get('id', '')) else ''
                    url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
                    rows_fac.append({
                        'Publication Title': str(r['title']),
                        'Year': int(r['year']) if pd.notna(r['year']) else '',
                        'Citations': int(r['times_cited']),
                        '_url': url
                    })
                _add_hyperlink_table(
                    doc, rows_fac,
                    col_headers=['Publication Title', 'Year', 'Citations'],
                    col_widths_pct=[75, 10, 15],
                    url_col_idx=0,
                    caption=TBL('Top publications co-authored with facility organisations, ranked by citations')
                )
            else:
                _add_body(doc, 'No facility co-authored publications found.', italic=True)
        else:
            _add_body(doc, '(Organisation type data not available.)', italic=True)
    except Exception as e:
        print(f'⚠️   Could not build Facility Co-authorship section: {e}')

    # ── Nonprofit Co-authorship ──────────────────────────────────────────────
    _add_heading(doc, 'Nonprofit Co-authorship', level=3)
    _add_body(doc, (
        'The following table shows publications co-authored with researchers in this '
        'report where at least one co-authoring organisation is classified as a Nonprofit '
        'in Dimensions. This includes charitable organisations, foundations, and '
        'non-governmental organisations.'
    ))
    try:
        if not cp.empty and 'research_org_types' in cp.columns:
            def _has_nonprofit(val):
                if isinstance(val, list):
                    return any('nonprofit' in str(t).lower() for t in val)
                return 'nonprofit' in str(val).lower()
            nonprofit_pubs = cp[cp['research_org_types'].apply(_has_nonprofit)].copy()
            if not nonprofit_pubs.empty:
                nonprofit_pubs['times_cited'] = pd.to_numeric(
                    nonprofit_pubs.get('times_cited', 0), errors='coerce').fillna(0)
                yr_nfp = (nonprofit_pubs.groupby('year').size().reset_index(name='Count')
                          .dropna(subset=['year']))
                yr_nfp['year'] = yr_nfp['year'].astype(int)
                if len(yr_nfp) > 1:
                    buf = chart_col(yr_nfp['year'].astype(str).tolist(),
                                    yr_nfp['Count'].tolist(),
                                    'Nonprofit Co-authored Publications by Year',
                                    xlabel='Year', ylabel='No. Publications', figsize=(6, 3))
                    _insert_image(doc, buf, 5.5,
                                  FIG('Nonprofit co-authored publications by year'))
                top_nfp = (nonprofit_pubs[['id', 'title', 'year', 'times_cited']]
                           .sort_values('times_cited', ascending=False)
                           .head(10).reset_index(drop=True))
                rows_nfp = []
                for _, r in top_nfp.iterrows():
                    pub_id = str(r['id']) if pd.notna(r.get('id', '')) else ''
                    url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
                    rows_nfp.append({
                        'Publication Title': str(r['title']),
                        'Year': int(r['year']) if pd.notna(r['year']) else '',
                        'Citations': int(r['times_cited']),
                        '_url': url
                    })
                _add_hyperlink_table(
                    doc, rows_nfp,
                    col_headers=['Publication Title', 'Year', 'Citations'],
                    col_widths_pct=[75, 10, 15],
                    url_col_idx=0,
                    caption=TBL('Top publications co-authored with nonprofit organisations, ranked by citations')
                )
            else:
                _add_body(doc, 'No nonprofit co-authored publications found.', italic=True)
        else:
            _add_body(doc, '(Organisation type data not available.)', italic=True)
    except Exception as e:
        print(f'⚠️   Could not build Nonprofit Co-authorship section: {e}')

    # ── Top Researchers by Impact ─────────────────────────────────────────────
    _build_top_researchers_section(doc, data, meta, tbl)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 3
    # ══════════════════════════════════════════════════════════════════════════
    _page_break(doc)
    _add_heading(doc, 'Part 3: Pathways to Impact', level=1)
    _add_body(doc, (
        f'The following metrics are sourced from the Dimensions database. Selections were based on '
        f'outputs from the {meta.get("faculty","department")} published during {meta.get("period","")}.'
    ))

    _add_heading(doc, 'Summary of Impact', level=3)
    _add_body(doc, (
        'This section provides an overview of research activity and performance. It includes total '
        'publication output and annual trends, alongside citation metrics and FCR.'
    ))
    _add_body(doc, (
        'Broader engagement is reflected through Altmetric attention, highlighting visibility beyond '
        'academic citations. The section also reports on Open Access and closed access outputs, '
        'enabling monitoring of accessibility and progress toward open research.'
    ))

    # ── Table 5 – Publication Indicators for Impact ───────────────────────────
    imp_metrics = [
        'Total Publications', 'Total Citations',
        'Publications with Altmetric Score',
        'Open Access Publications (oa_all: gold, hybrid, bronze, green)',
        'Closed Access Publications'
    ]
    rows_imp = []
    for m in imp_metrics:
        row = ps[ps['Metric'] == m]
        if not row.empty:
            r = {'Metric': m}
            r['Overall'] = row['Overall'].values[0]
            for c in yr_c:
                r[c] = row[c].values[0] if c in row.columns else ''
            rows_imp.append(r)
    if rows_imp:
        imp_df = pd.DataFrame(rows_imp)
        for col in ['Overall'] + list(yr_c):
            imp_df[col] = imp_df[col].apply(
                lambda v: f'{int(float(v)):,}' if pd.notna(v) and str(v) != '' else '')
        # Shorten long metric names so numbers have room to fit
        metric_short = {
            'Total Publications':                                              'Total Publications',
            'Total Citations':                                                 'Total Citations',
            'Publications with Altmetric Score':                               'With Altmetric Score',
            'Open Access Publications (oa_all: gold, hybrid, bronze, green)': 'Open Access',
            'Closed Access Publications':                                      'Closed Access',
        }
        imp_df['Metric'] = imp_df['Metric'].map(lambda x: metric_short.get(x, x))
        imp_df = imp_df.rename(columns={'Overall': 'All'})
        yr_cols_display = ['All'] + list(yr_c)
        yr_count = len(yr_cols_display)
        # Each numeric col needs ~10% to fit "105,142"; metric gets the rest
        num_pct  = 10
        metric_pct = 100 - num_pct * yr_count
        col_pcts = [metric_pct] + [num_pct] * yr_count
        col_pcts[0] += 100 - sum(col_pcts)
        _add_table(doc, imp_df,
                   caption=TBL('Publication Indicators for Academic and Societal Impact'),
                   col_widths_pct=col_pcts)

    # ── Policy Citations ──────────────────────────────────────────────────────
    _add_heading(doc, 'Policy Citations', level=3)

    # Publisher summary table
    pol_df = data['policy'][['Publisher','Policy Documents','Publications Cited']].copy()
    pol_df = (pol_df.dropna(subset=['Publications Cited'])
                    .sort_values('Publications Cited', ascending=False)
                    .head(10).reset_index(drop=True))
    pol_df.columns = ['Publisher','Policy Docs','Pubs Cited']
    _add_table(doc, pol_df,
               caption=TBL('Top Policy Publishers by Publications Cited'),
               col_widths_pct=[65, 18, 17])

    # Top publications cited by policy (hyperlinked titles)
    pol_pubs = (cp[cp['policy_citations'] > 0]
                [['id','title','doi','year','policy_citations']]
                .sort_values('policy_citations', ascending=False)
                .head(10).reset_index(drop=True))
    if not pol_pubs.empty:
        rows_pol = []
        for _, r in pol_pubs.iterrows():
            pub_id = str(r['id']) if pd.notna(r.get('id','')) else ''
            url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
            rows_pol.append({
                'Publication Title': str(r['title']),
                'Year': int(r['year']) if pd.notna(r['year']) else '',
                'Policy Citations': int(r['policy_citations']),
                '_url': url
            })
        _add_hyperlink_table(
            doc, rows_pol,
            col_headers=['Publication Title','Year','Policy Citations'],
            col_widths_pct=[68, 10, 22],
            url_col_idx=0,
            caption=TBL('Top publications by Policy Citations')
        )

    # ── Patent Citations ──────────────────────────────────────────────────────
    _add_heading(doc, 'Patent Citations', level=3)

    pat_df = data['patents'][['Assignee','Country','Patent Citations','Publications Cited']].copy()
    pat_df = (pat_df.dropna(subset=['Assignee'])
                    .query('Assignee.str.strip() != ""', engine='python')
                    .loc[lambda d: ~d['Assignee'].str.startswith(('─','—','─','—'), na=False)]
                    .head(10).reset_index(drop=True))
    pat_df.columns = ['Assignee','Country','Patent Citations','Pubs Cited']
    _add_table(doc, pat_df,
               caption=TBL('Top Patent Assignees citing School Publications'),
               col_widths_pct=[48, 20, 16, 16])

    pat_pubs = (cp[cp['patent_citations'] > 0]
                [['id','title','doi','year','patent_citations']]
                .sort_values('patent_citations', ascending=False)
                .head(10).reset_index(drop=True))
    if not pat_pubs.empty:
        rows_pat = []
        for _, r in pat_pubs.iterrows():
            pub_id = str(r['id']) if pd.notna(r.get('id','')) else ''
            url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
            rows_pat.append({
                'Publication Title': str(r['title']),
                'Year': int(r['year']) if pd.notna(r['year']) else '',
                'Patent Citations': int(r['patent_citations']),
                '_url': url
            })
        _add_hyperlink_table(
            doc, rows_pat,
            col_headers=['Publication Title','Year','Patent Citations'],
            col_widths_pct=[68, 10, 22],
            url_col_idx=0,
            caption=TBL('Top publications by Patent Citations')
        )

    # ── Clinical Trial Citations ──────────────────────────────────────────────
    _add_heading(doc, 'Clinical Trial Citations', level=3)

    ct_df = data['clinical'][['Sponsor / Collaborator','Country','Trial Citations','Publications Cited']].copy()
    ct_df = (ct_df.dropna(subset=['Sponsor / Collaborator'])
                  .query('`Sponsor / Collaborator`.str.strip() != ""', engine='python')
                  .loc[lambda d: ~d['Sponsor / Collaborator'].str.startswith(('─','—','─','—'), na=False)]
                  .head(10).reset_index(drop=True))
    ct_df.columns = ['Sponsor / Collaborator','Country','Trial Citations','Pubs Cited']
    _add_table(doc, ct_df,
               caption=TBL('Top Clinical Trial Sponsors/Collaborators citing School Publications'),
               col_widths_pct=[46, 20, 17, 17])

    ct_pubs = (cp[cp['clinical_citations'] > 0]
               [['id','title','doi','year','clinical_citations']]
               .sort_values('clinical_citations', ascending=False)
               .head(10).reset_index(drop=True))
    if not ct_pubs.empty:
        rows_ct = []
        for _, r in ct_pubs.iterrows():
            pub_id = str(r['id']) if pd.notna(r.get('id','')) else ''
            url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
            rows_ct.append({
                'Publication Title': str(r['title']),
                'Year': int(r['year']) if pd.notna(r['year']) else '',
                'Clinical Citations': int(r['clinical_citations']),
                '_url': url
            })
        _add_hyperlink_table(
            doc, rows_ct,
            col_headers=['Publication Title','Year','Clinical Citations'],
            col_widths_pct=[68, 10, 22],
            url_col_idx=0,
            caption=TBL('Top publications by Clinical Trial Citations')
        )

    # ── Altmetric Attention ───────────────────────────────────────────────────
    _add_heading(doc, 'Most Alternative Research Attention', level=3)
    _add_body(doc, (
        'The following metrics for research attention are sourced from the Altmetric Explorer and '
        'cover a range of alternative research attention and impact measures, including mentions in '
        'news, policy and social media. Altmetric data can only track publications with a unique '
        'identifier such as a DOI, when that identifier is used in an online mention.'
    ))
    _add_body(doc, 'Below are the top 10 research outputs by Altmetric Attention Score.')

    if 'altmetric' in cp.columns:
        alt = (cp[cp['altmetric'].notna() & (cp['altmetric'] > 0)]
               [['id','title','doi','year','type','altmetric']]
               .sort_values('altmetric', ascending=False)
               .head(10).reset_index(drop=True))
        rows_alt = []
        for _, r in alt.iterrows():
            doi = str(r['doi']) if pd.notna(r['doi']) else ''
            pub_id = str(r['id']) if pd.notna(r.get('id','')) else ''
            url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
            rows_alt.append({
                'Publication Title': str(r['title']),
                'DOI': doi,
                'Year': int(r['year']) if pd.notna(r['year']) else '',
                'Type': str(r['type']) if pd.notna(r['type']) else '',
                'Altmetric Score': int(r['altmetric']) if pd.notna(r['altmetric']) else '',
                '_url': url
            })
        _add_hyperlink_table(
            doc, rows_alt,
            col_headers=['Publication Title','DOI','Year','Type','Altmetric Score'],
            col_widths_pct=[42, 24, 8, 10, 16],
            url_col_idx=0,
            caption=TBL('Top 10 Publications by Altmetric Attention Score')
        )

    # ── Corporate Citations top pubs ──────────────────────────────────────────
    if 'corporate_citations' in cp.columns:
        corp_pubs = (cp[pd.to_numeric(cp['corporate_citations'], errors='coerce') > 0]
                     [['id','title','doi','year','corporate_citations']]
                     .copy())
        corp_pubs['corporate_citations'] = pd.to_numeric(corp_pubs['corporate_citations'], errors='coerce')
        corp_pubs = corp_pubs.sort_values('corporate_citations', ascending=False).head(10).reset_index(drop=True)
        if not corp_pubs.empty:
            _add_heading(doc, 'Corporate Citations', level=3)
            _add_body(doc, 'The top 10 publications by corporate citations are shown below.').paragraph_format.keep_with_next = True
            rows_corp = []
            for _, r in corp_pubs.iterrows():
                pub_id = str(r['id']) if pd.notna(r.get('id','')) else ''
                url = f'https://app.dimensions.ai/details/publication/{pub_id}' if pub_id and pub_id != 'nan' else ''
                rows_corp.append({
                    'Publication Title': str(r['title']),
                    'Year': int(r['year']) if pd.notna(r['year']) else '',
                    'Corporate Citations': int(r['corporate_citations']),
                    '_url': url
                })
            _add_hyperlink_table(
                doc, rows_corp,
                col_headers=['Publication Title', 'Year', 'Corporate Citations'],
                col_widths_pct=[75, 10, 15],
                url_col_idx=0,
                caption=TBL('Top 10 publications by corporate citations')
            )

    # ══════════════════════════════════════════════════════════════════════════
    # PART 4 – UNIVERSITY COMPARISON BY FIELD OF RESEARCH
    # ══════════════════════════════════════════════════════════════════════════
    _page_break(doc)
    _add_heading(doc, 'Part 4: University Comparison by Field of Research', level=1)
    _add_body(doc, (
        'Comparisons of research performance between institutions cannot be carried out at the '
        'department or school level, as equivalent lists of researchers from other institutions '
        'are not available for comparison. Instead, comparisons are made using all publications '
        'produced by each institution within specific Fields of Research.'
    ))

    bench = data.get('bench', {})
    if bench:
        _build_university_comparison(doc, bench, meta, fig, tbl)
    else:
        _add_body(doc, '(No benchmarking file detected. Place the institutional_benchmarking_report xlsx '
                        'in the same folder as the metrics export and re-run to populate this section.)',
                  italic=True)

    # ── Appendix ──────────────────────────────────────────────────────────────
    _build_appendix(doc, data)


# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX
# ─────────────────────────────────────────────────────────────────────────────

def _build_appendix(doc, data):
    """Appendix: Methodology & Report Information from both source files."""

    def _get(df, label):
        for _, r in df.iterrows():
            if str(r.iloc[0]).strip() == label:
                v = r.iloc[1]
                return '' if pd.isna(v) else str(v).strip()
        return ''

    def _tbl(doc, rows, col_widths_pct=None):
        """Two-column label/value table styled to match the rest of the report."""
        col_widths_pct = col_widths_pct or [33, 67]
        total = 9026
        widths = [int(total * p / 100) for p in col_widths_pct]
        _add_table(doc, rows, col_widths_pct=col_widths_pct)

    doc.add_page_break()
    _add_heading(doc, 'Appendix: Methodology & Report Information', level=1)

    # ── A. Researcher Metrics ─────────────────────────────────────────────────
    _add_heading(doc, 'A. Researcher Metrics Report', level=3)
    _add_body(doc,
        'The following report details and methodology notes are drawn from the Report '
        'Information tab of the Researcher Metrics Report Excel file. '
        'Data is sourced from the Dimensions database (dimensions.ai).')

    rm = data.get('report_info', pd.DataFrame())
    if rm.empty:
        _add_body(doc, '(Report information not available.)', italic=True)
        return

    # Report Details
    detail_labels = [
        'Date Generated', 'Institution', 'Faculty / Department',
        'Reporting Period', 'Number of Researchers', 'Total Publications', 'Data Source',
    ]
    details = [(lbl, _get(rm, lbl)) for lbl in detail_labels if _get(rm, lbl)]
    if details:
        _add_body(doc, 'Report Details', bold=True)
        df_details = pd.DataFrame(details, columns=['Field', 'Value'])
        _add_table(doc, df_details, caption=None, col_widths_pct=[33, 67])

    # Methodology Notes
    method_labels = [
        'Patent Citations', 'Policy Citations', 'Corporate Citations',
        'Field Citation Ratio (FCR)', 'Open Access', 'Organisation Types',
        'Collaboration Type', 'Clinical Trial Sponsors', 'Researcher Citations',
    ]
    methods = [(lbl, _get(rm, lbl)) for lbl in method_labels if _get(rm, lbl)]
    if methods:
        _add_body(doc, 'Methodology Notes', bold=True)
        df_methods = pd.DataFrame(methods, columns=['Metric', 'Notes'])
        _add_table(doc, df_methods, caption=None, col_widths_pct=[25, 75])

    # Tab Descriptions
    tab_start = None
    for i, r in rm.iterrows():
        if str(r.iloc[0]).strip() == 'Tab Name':
            tab_start = i + 1
            break
    if tab_start:
        tabs = []
        for i in range(tab_start, len(rm)):
            name = str(rm.iloc[i, 0]).strip()
            desc = str(rm.iloc[i, 1]).strip() if pd.notna(rm.iloc[i, 1]) else ''
            if name and name != 'nan' and desc and desc != 'nan':
                tabs.append((name, desc))
        if tabs:
            _add_body(doc, 'Data Sheet Descriptions', bold=True)
            df_tabs = pd.DataFrame(tabs, columns=['Sheet', 'Description'])
            _add_table(doc, df_tabs, caption=None, col_widths_pct=[25, 75])

    # Footer note
    for _, r in rm.iterrows():
        v = str(r.iloc[0]).strip()
        if v.startswith('Generated using'):
            _add_body(doc, v, italic=True)
            break

    # ── B. Institutional Benchmarking ─────────────────────────────────────────
    bench = data.get('bench', {})
    bench_path = bench.get('_path') if bench else None
    if not bench_path:
        return

    try:
        bm = pd.read_excel(bench_path, sheet_name='Info', header=None)
    except Exception:
        return

    _add_heading(doc, 'B. Institutional Benchmarking Report', level=3)
    _add_body(doc,
        'The following methodology notes are drawn from the Info tab of the Institutional '
        'Benchmarking Report Excel file.')

    # Parameters
    institutions = _get(bm, 'Institutions').replace('\n', '; ')
    bm_params = [
        ('Year Range',    _get(bm, 'Year Range')),
        ('FoR Codes',     _get(bm, 'FoR Codes')),
        ('Institutions',  institutions),
        ('Data Source',   _get(bm, 'Data Source')),
    ]
    bm_params = [(k, v) for k, v in bm_params if v]
    if bm_params:
        _add_body(doc, 'Report Parameters', bold=True)
        df_params = pd.DataFrame(bm_params, columns=['Parameter', 'Value'])
        _add_table(doc, df_params, caption=None, col_widths_pct=[25, 75])

    # Tab Descriptions
    # Sheet descriptions — overridden here to reflect global-only methodology
    bm_tabs = [
        ('All Publications',
         _get(bm, 'All Publications') or
         'Full list of all publications retrieved for all institutions across all FoR codes entered. '
         'Includes publication ID, title, year, citations, FCR, and flags indicating which FoR code(s) '
         'each paper belongs to.'),
        ('Summary (FoR code)',
         'One row per institution showing aggregate metrics: total publications, total citations, '
         'citations per paper, publications with FCR, FCR geo mean, and % publications with FCR.'),
        ('Top Cited (FoR code)',
         "One table per year showing each institution's count and percentage of publications exceeding "
         'the global Top 1%, 5%, and 10% citation thresholds. ALL GLOBAL row shows the total global '
         'publication count and the percentage of global publications exceeding each threshold '
         '(computed via API count queries).'),
    ]
    bm_tabs = [(k, v) for k, v in bm_tabs if v]
    if bm_tabs:
        _add_body(doc, 'Data Sheet Descriptions', bold=True)
        df_btabs = pd.DataFrame(bm_tabs, columns=['Sheet', 'Description'])
        _add_table(doc, df_btabs, caption=None, col_widths_pct=[25, 75])

    # Top Cited Methodology
    # Note: 'Two benchmark pools' row is intentionally excluded — the report uses
    # global thresholds only. The row is replaced with a single 'Benchmark pool' entry.
    tc_labels = [
        'Step 1 — Citation threshold', 'Step 2 — Institutional count',
        'ALL GLOBAL percentages', 'Important caveat', 'FCR (Field Citation Ratio)',
    ]
    tc_methods = [(lbl, _get(bm, lbl)) for lbl in tc_labels if _get(bm, lbl)]
    # Insert single global benchmark pool row after 'ALL GLOBAL percentages'
    benchmark_pool_row = (
        'Benchmark pool',
        'Global — thresholds derived from all publications worldwide. '
        'All percentile figures in this report reflect international comparison benchmarks.'
    )
    new_tc_methods = []
    for row in tc_methods:
        new_tc_methods.append(row)
        if row[0] == 'ALL GLOBAL percentages':
            new_tc_methods.append(benchmark_pool_row)
    if not any(r[0] == 'Benchmark pool' for r in new_tc_methods):
        new_tc_methods.append(benchmark_pool_row)
    if new_tc_methods:
        _add_body(doc, 'Top Cited Percentile Methodology', bold=True)
        df_tc = pd.DataFrame(new_tc_methods, columns=['Step', 'Description'])
        _add_table(doc, df_tc, caption=None, col_widths_pct=[25, 75])


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    import glob

    def _find_metrics_file(hint=None):
        """Return the metrics xlsx path, auto-detecting if needed."""
        # If a path was given and it exists, use it directly
        if hint and os.path.exists(hint):
            return hint
        # If a path was given but not found, try glob in same dir for versioned copies
        if hint:
            base_dir = os.path.dirname(os.path.abspath(hint))
            stem = os.path.splitext(os.path.basename(hint))[0]
            # e.g. Researcher_Metrics_Report_2020_2025(2).xlsx
            candidates = glob.glob(os.path.join(base_dir, stem + '*.xlsx'))
            if candidates:
                candidates.sort(key=os.path.getmtime, reverse=True)
                return candidates[0]
        # No hint or still not found — look in current directory for any metrics file
        patterns = [
            'Researcher_Metrics_Report*.xlsx',
            '*Researcher*Metrics*.xlsx',
            '*researcher*metrics*.xlsx',
        ]
        for pat in patterns:
            matches = glob.glob(pat)
            if matches:
                matches.sort(key=os.path.getmtime, reverse=True)
                return matches[0]
        return None

    hint      = sys.argv[1] if len(sys.argv) >= 2 else None
    xlsx_path = _find_metrics_file(hint)

    if not xlsx_path:
        print('❌  Could not find the Researcher Metrics Excel file.')
        print('    Place it in the same folder as generate_report.py and try again.')
        sys.exit(1)

    if hint and xlsx_path != hint:
        print(f'⚠️   "{hint}" not found — using "{os.path.basename(xlsx_path)}" instead.')

    template_path = sys.argv[2] if len(sys.argv) > 2 else \
        'University_of_New_South_Wales_Departmental_Report_2026.docx'
    out_name      = sys.argv[3] if len(sys.argv) > 3 else \
        f'Departmental_Report_{datetime.date.today().isoformat()}.docx'

    print(f'📖  Reading data from: {xlsx_path}')
    data = load_excel(xlsx_path)
    meta = get_report_meta(data)

    print(f'📝  Faculty  : {meta.get("faculty","—")}')
    print(f'📅  Period   : {meta.get("period","—")}')
    print(f'📅  Date     : {meta.get("date","—")} (AEDT)')

    # Fresh document — template namespace pollution breaks fixed-width tables in LibreOffice
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Twips(11906)
    sec.page_height   = Twips(16838)
    sec.left_margin   = Twips(1440)
    sec.right_margin  = Twips(1440)
    sec.top_margin    = Twips(1440)
    sec.bottom_margin = Twips(1440)

    if os.path.exists(template_path):
        print(f'📄  Template found (A4 page settings applied)')

    _setup_styles(doc)
    print('🔨  Building report...')
    section_cover(doc, meta)
    build_report(doc, data, meta)

    doc.save(out_name)
    print(f'\n✅  Report saved to: {out_name}')
    return out_name

    print(f'📝  Faculty  : {meta.get("faculty","—")}')
    print(f'📅  Period   : {meta.get("period","—")}')
    print(f'📅  Date     : {meta.get("date","—")} (AEDT)')

    # Fresh document — template namespace pollution breaks fixed-width tables in LibreOffice
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Twips(11906)
    sec.page_height   = Twips(16838)
    sec.left_margin   = Twips(1440)
    sec.right_margin  = Twips(1440)
    sec.top_margin    = Twips(1440)
    sec.bottom_margin = Twips(1440)

    if os.path.exists(template_path):
        print(f'📄  Template found (A4 page settings applied)')

    _setup_styles(doc)
    print('🔨  Building report...')
    section_cover(doc, meta)
    build_report(doc, data, meta)

    doc.save(out_name)
    print(f'\n✅  Report saved to: {out_name}')
    return out_name


if __name__ == '__main__':
    main()
